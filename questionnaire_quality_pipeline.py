#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Questionnaire quality and measurement-structure assessment pipeline for Likert-scale Excel data.

This tool provides reliability evidence and preliminary validity evidence based on
response processes available in the input file and the questionnaire's internal
structure. It does not claim that a questionnaire has achieved comprehensive
validity from EFA, AVE, HTMT, or internal-consistency coefficients alone.

Outputs
-------
- Questionnaire_quality_report.docx: concise tables and narrative interpretation
- Questionnaire_quality_report.html: full interactive-style visual report
- Questionnaire_quality_details.xlsx: complete diagnostics and alternative solutions
- figures/*.png: scree plot, loading heatmaps, correlations, reliability, stability
- CFA_lavaan.R / ESEM_lavaan.R: follow-up R scripts when confirmatory work is feasible
- paper_ready_summary.txt: multiple manuscript-ready reporting scenarios plus references

Automatic item deletion is conservative. Protected content-essential items are never
automatically deleted, small samples do not trigger permanent deletion by default,
and planned dimensions are retained unless there is strong counterevidence.
"""

from __future__ import annotations

import argparse
import base64
from copy import copy
import io
import json
import math
import os
import shutil
import subprocess
import sys
import textwrap
import traceback
import warnings
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Iterable, Optional

import matplotlib.pyplot as plt

# Prefer Traditional Chinese-capable fonts when available; fall back safely on other systems.
plt.rcParams["font.sans-serif"] = [
    "Noto Sans CJK TC", "Microsoft JhengHei", "PingFang TC",
    "Arial Unicode MS", "DejaVu Sans"
]
plt.rcParams["axes.unicode_minus"] = False

import numpy as np
import pandas as pd
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from jinja2 import Template
from scipy.linalg import eigh
from scipy.optimize import linear_sum_assignment, minimize_scalar
from scipy.stats import chi2, multivariate_normal, norm
from statsmodels.multivariate.factor import Factor
from statsmodels.stats.correlation_tools import corr_nearest

warnings.filterwarnings("ignore", category=RuntimeWarning)


# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------


@dataclass
class Thresholds:
    loading_minimum: float = 0.32
    loading_preferred: float = 0.40
    cross_loading: float = 0.30
    loading_gap: float = 0.20
    communality_severe: float = 0.20
    communality_warning: float = 0.40
    msa_minimum: float = 0.50
    factor_correlation_oblique: float = 0.20
    redundancy_correlation: float = 0.85
    htmt_strict: float = 0.85
    htmt_lenient: float = 0.90


@dataclass
class Config:
    data_sheet: str = "responses"
    codebook_sheet: str = "codebook"
    id_columns: list[str] = field(default_factory=lambda: ["ID", "id", "編號"])
    exclude_columns: list[str] = field(default_factory=list)
    expected_factors: Optional[int] = None
    likert_min: Optional[float] = None
    likert_max: Optional[float] = None
    missing_row_exclusion: float = 0.50
    missing_item_warning: float = 0.10
    straightline_sd: float = 0.05
    longstring_fraction: float = 0.80
    max_delete_fraction: float = 0.35
    max_delete_iterations: int = 20
    minimum_items_per_factor: int = 3
    random_seed: int = 20260730
    pa_iterations_pearson: int = 300
    pa_iterations_polychoric: int = 40
    bootstrap_iterations: int = 200
    bootstrap_polychoric_iterations: int = 50
    run_confirmatory_if_feasible: bool = True
    run_r_if_available: bool = False
    prioritize_planned_dimensions: bool = True
    report_alternative_factor_solution: bool = True
    small_sample_threshold: int = 50
    preserve_planned_below_n: int = 50
    disable_auto_deletion_below_n: int = 50
    planned_factor_correlation_critical: float = 0.90
    planned_rmsr_critical: float = 0.10
    planned_problem_fraction_critical: float = 0.30
    merge_assignment_minimum: float = 0.60
    split_assignment_minimum: float = 0.30
    output_language: str = "zh-TW"
    thresholds: Thresholds = field(default_factory=Thresholds)

    @staticmethod
    def from_yaml(path: Optional[str]) -> "Config":
        cfg = Config()
        if not path:
            return cfg
        raw = yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}
        t_raw = raw.pop("thresholds", {}) or {}
        for key, val in raw.items():
            if hasattr(cfg, key):
                setattr(cfg, key, val)
        for key, val in t_raw.items():
            if hasattr(cfg.thresholds, key):
                setattr(cfg.thresholds, key, val)
        return cfg


# -----------------------------------------------------------------------------
# Small helpers
# -----------------------------------------------------------------------------


def safe_float(x: Any, default: float = np.nan) -> float:
    try:
        return float(x)
    except Exception:
        return default


def fmt(x: Any, digits: int = 3, na: str = "NA") -> str:
    try:
        if x is None or not np.isfinite(float(x)):
            return na
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def boolish(x: Any) -> bool:
    if isinstance(x, bool):
        return x
    if pd.isna(x):
        return False
    return str(x).strip().lower() in {"1", "true", "yes", "y", "是", "反向"}


def sanitize_sheet_name(name: str) -> str:
    bad = "[]:*?/\\"
    for c in bad:
        name = name.replace(c, "_")
    return name[:31]


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def offdiag_values(a: np.ndarray) -> np.ndarray:
    mask = ~np.eye(a.shape[0], dtype=bool)
    return a[mask]


def nearest_correlation(r: np.ndarray) -> tuple[np.ndarray, dict[str, float]]:
    r = np.asarray(r, dtype=float)
    r = (r + r.T) / 2
    np.fill_diagonal(r, 1.0)
    eig_before = np.linalg.eigvalsh(r)
    need_fix = np.nanmin(eig_before) <= 1e-8 or not np.all(np.isfinite(r))
    r0 = np.nan_to_num(r, nan=0.0, posinf=0.99, neginf=-0.99)
    np.fill_diagonal(r0, 1.0)
    if need_fix:
        fixed = corr_nearest(r0, threshold=1e-6, n_fact=100)
    else:
        fixed = r0
    fixed = np.asarray(fixed)
    fixed = (fixed + fixed.T) / 2
    np.fill_diagonal(fixed, 1.0)
    diff = fixed - r0
    return fixed, {
        "minimum_eigenvalue_before": float(np.min(eig_before)),
        "minimum_eigenvalue_after": float(np.min(np.linalg.eigvalsh(fixed))),
        "rms_adjustment": float(np.sqrt(np.mean(diff**2))),
        "maximum_absolute_adjustment": float(np.max(np.abs(diff))),
        "was_adjusted": bool(need_fix),
    }


def longstring_length(row: pd.Series) -> int:
    values = [v for v in row.tolist() if not pd.isna(v)]
    if not values:
        return 0
    best = cur = 1
    for i in range(1, len(values)):
        if values[i] == values[i - 1]:
            cur += 1
            best = max(best, cur)
        else:
            cur = 1
    return best


def reorder_square(matrix: np.ndarray, order: list[int]) -> np.ndarray:
    return matrix[np.ix_(order, order)]


# -----------------------------------------------------------------------------
# Input and codebook
# -----------------------------------------------------------------------------


def read_input(excel_path: str, cfg: Config) -> tuple[pd.DataFrame, pd.DataFrame, list[str], list[str]]:
    xls = pd.ExcelFile(excel_path)
    if cfg.data_sheet not in xls.sheet_names:
        data_sheet = xls.sheet_names[0]
    else:
        data_sheet = cfg.data_sheet
    data = pd.read_excel(excel_path, sheet_name=data_sheet)
    data.columns = [str(c).strip() for c in data.columns]

    if cfg.codebook_sheet in xls.sheet_names:
        codebook = pd.read_excel(excel_path, sheet_name=cfg.codebook_sheet)
        codebook.columns = [str(c).strip() for c in codebook.columns]
    else:
        codebook = pd.DataFrame()

    messages: list[str] = []
    if not codebook.empty and "item" in codebook.columns:
        if "include" in codebook.columns:
            cb = codebook[codebook["include"].map(lambda x: True if pd.isna(x) else boolish(x))].copy()
        else:
            cb = codebook.copy()
        items = [str(x).strip() for x in cb["item"].dropna().tolist()]
        missing_items = [x for x in items if x not in data.columns]
        if missing_items:
            messages.append(f"Codebook中有{len(missing_items)}個題目欄位未出現在資料表：{', '.join(missing_items)}")
        items = [x for x in items if x in data.columns]
    else:
        excluded = set(cfg.exclude_columns) | set(cfg.id_columns)
        items = []
        for col in data.columns:
            if col in excluded:
                continue
            numeric = pd.to_numeric(data[col], errors="coerce")
            nunique = numeric.nunique(dropna=True)
            if numeric.notna().mean() >= 0.80 and 2 <= nunique <= 10:
                items.append(col)
        codebook = pd.DataFrame({
            "item": items,
            "item_text": items,
            "expected_factor": np.nan,
            "reverse": False,
            "include": True,
        })
        messages.append("找不到codebook工作表，因此以『數值型且有2至10個反應類別』自動辨識題目。正式分析建議補上codebook。")

    if len(items) < 4:
        raise ValueError(f"可分析題目只有{len(items)}題；EFA至少需要更多題目。請檢查codebook或欄位格式。")

    id_cols = [c for c in cfg.id_columns if c in data.columns]
    return data, codebook, items, messages + [f"讀取資料工作表：{data_sheet}"]


def prepare_items(
    raw_data: pd.DataFrame,
    codebook: pd.DataFrame,
    items: list[str],
    cfg: Config,
) -> tuple[pd.DataFrame, pd.DataFrame, list[str]]:
    data = raw_data.copy()
    messages: list[str] = []
    for item in items:
        data[item] = pd.to_numeric(data[item], errors="coerce")

    cb = codebook.copy()
    if "item" not in cb.columns:
        cb["item"] = items
    cb["item"] = cb["item"].astype(str).str.strip()
    cb = cb[cb["item"].isin(items)].drop_duplicates("item").copy()
    cb = cb.set_index("item").reindex(items).reset_index()
    for col, default in [
        ("item_text", ""),
        ("expected_factor", np.nan),
        ("reverse", False),
        ("likert_min", np.nan),
        ("likert_max", np.nan),
        ("protect", False),
        ("content_note", ""),
    ]:
        if col not in cb.columns:
            cb[col] = default

    for _, row in cb.iterrows():
        item = row["item"]
        if not boolish(row.get("reverse", False)):
            continue
        low = safe_float(row.get("likert_min"), cfg.likert_min if cfg.likert_min is not None else np.nan)
        high = safe_float(row.get("likert_max"), cfg.likert_max if cfg.likert_max is not None else np.nan)
        if not np.isfinite(low):
            low = data[item].min(skipna=True)
        if not np.isfinite(high):
            high = data[item].max(skipna=True)
        if np.isfinite(low) and np.isfinite(high):
            data[item] = low + high - data[item]
            messages.append(f"反向計分：{item}（範圍 {fmt(low,0)}–{fmt(high,0)}）")
        else:
            messages.append(f"無法反向計分：{item}，因量尺上下限無法判定。")

    return data, cb, messages


# -----------------------------------------------------------------------------
# Data checks
# -----------------------------------------------------------------------------


def data_checks(data: pd.DataFrame, items: list[str], cfg: Config) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    x = data[items].copy()
    n0 = len(x)
    row_missing = x.isna().mean(axis=1)
    row_sd = x.std(axis=1, skipna=True)
    row_longstring = x.apply(longstring_length, axis=1)
    flagged_row = (
        (row_missing > cfg.missing_row_exclusion)
        | (row_sd.fillna(0) <= cfg.straightline_sd)
        | (row_longstring >= np.ceil(len(items) * cfg.longstring_fraction))
    )
    row_diag = pd.DataFrame({
        "row_index": x.index,
        "missing_fraction": row_missing,
        "within_person_sd": row_sd,
        "longstring": row_longstring,
        "flagged": flagged_row,
        "reason": [
            "; ".join([
                "過多遺漏" if row_missing.loc[i] > cfg.missing_row_exclusion else "",
                "作答變異近零" if safe_float(row_sd.loc[i], 0) <= cfg.straightline_sd else "",
                "長串相同作答" if row_longstring.loc[i] >= np.ceil(len(items) * cfg.longstring_fraction) else "",
            ]).strip("; ").replace("; ;", ";")
            for i in x.index
        ],
    })

    # Only rows with more than half missing are excluded automatically. Straight-lining is sensitivity information.
    auto_exclude = row_missing > cfg.missing_row_exclusion
    clean = x.loc[~auto_exclude].copy()

    rows = []
    freq_rows = []
    for item in items:
        s = clean[item]
        vc = s.value_counts(dropna=False).sort_index()
        nonmiss = s.dropna()
        probs = nonmiss.value_counts(normalize=True) if len(nonmiss) else pd.Series(dtype=float)
        floor = float(probs.loc[nonmiss.min()]) if len(nonmiss) and nonmiss.min() in probs.index else np.nan
        ceiling = float(probs.loc[nonmiss.max()]) if len(nonmiss) and nonmiss.max() in probs.index else np.nan
        rows.append({
            "item": item,
            "n_valid": int(s.notna().sum()),
            "missing_n": int(s.isna().sum()),
            "missing_fraction": float(s.isna().mean()),
            "mean": float(nonmiss.mean()) if len(nonmiss) else np.nan,
            "sd": float(nonmiss.std(ddof=1)) if len(nonmiss) > 1 else np.nan,
            "minimum": float(nonmiss.min()) if len(nonmiss) else np.nan,
            "maximum": float(nonmiss.max()) if len(nonmiss) else np.nan,
            "unique_categories": int(nonmiss.nunique()),
            "skewness": float(nonmiss.skew()) if len(nonmiss) > 2 else np.nan,
            "excess_kurtosis": float(nonmiss.kurt()) if len(nonmiss) > 3 else np.nan,
            "floor_fraction": floor,
            "ceiling_fraction": ceiling,
            "sparse_category_fraction": float((nonmiss.value_counts() < 5).mean()) if len(nonmiss) else np.nan,
            "zero_variance": bool(nonmiss.nunique() <= 1),
        })
        for category, count in vc.items():
            freq_rows.append({
                "item": item,
                "category": "Missing" if pd.isna(category) else category,
                "count": int(count),
                "percentage": float(count / len(s)) if len(s) else np.nan,
            })
    item_diag = pd.DataFrame(rows)
    frequencies = pd.DataFrame(freq_rows)

    summary = {
        "n_original": int(n0),
        "n_analysis": int(len(clean)),
        "rows_auto_excluded": int(auto_exclude.sum()),
        "rows_flagged_total": int(flagged_row.sum()),
        "n_items": len(items),
        "overall_missing_fraction": float(clean.isna().mean().mean()),
        "items_missing_over_warning": int((item_diag["missing_fraction"] > cfg.missing_item_warning).sum()),
        "zero_variance_items": item_diag.loc[item_diag["zero_variance"], "item"].tolist(),
        "median_categories": float(item_diag["unique_categories"].median()),
        "maximum_categories": int(item_diag["unique_categories"].max()),
        "ordinal_likert_like": bool((item_diag["unique_categories"] <= 7).mean() >= 0.80),
        "sparse_item_fraction": float((item_diag["sparse_category_fraction"] > 0.25).mean()),
        "severe_skew_items": int((item_diag["skewness"].abs() > 2).sum()),
    }
    return clean, item_diag, frequencies, row_diag, summary


# -----------------------------------------------------------------------------
# Correlation matrices
# -----------------------------------------------------------------------------


def pearson_correlation(data: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, Any]]:
    r = data.corr(method="pearson", min_periods=max(5, int(len(data) * 0.5)))
    fixed, diag = nearest_correlation(r.values)
    return pd.DataFrame(fixed, index=data.columns, columns=data.columns), diag


def ordinal_thresholds(x: np.ndarray, clip: float = 1e-5) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    valid = x[~np.isnan(x)]
    categories, counts = np.unique(valid, return_counts=True)
    if len(categories) < 2:
        raise ValueError("變項少於兩個反應類別")
    cum = np.cumsum(counts)[:-1] / counts.sum()
    cum = np.clip(cum, clip, 1 - clip)
    thresholds = np.concatenate(([-np.inf], norm.ppf(cum), [np.inf]))
    mapping = {cat: idx for idx, cat in enumerate(categories)}
    codes = np.array([mapping.get(v, -1) if not np.isnan(v) else -1 for v in x], dtype=int)
    return categories, thresholds, codes


def rectangle_probability(lower: np.ndarray, upper: np.ndarray, rho: float) -> float:
    cov = np.array([[1.0, rho], [rho, 1.0]])
    p = multivariate_normal.cdf(
        upper,
        mean=np.zeros(2),
        cov=cov,
        lower_limit=lower,
        maxpts=20000,
        abseps=1e-6,
        releps=1e-6,
    )
    return max(float(p), 1e-12)


def polychoric_pair(x: np.ndarray, y: np.ndarray) -> tuple[float, str]:
    mask = ~(np.isnan(x) | np.isnan(y))
    x = x[mask].astype(float)
    y = y[mask].astype(float)
    if len(x) < 10:
        return np.nan, "pairwise_n<10"
    try:
        cats_x, tx, cx = ordinal_thresholds(x)
        cats_y, ty, cy = ordinal_thresholds(y)
        table = np.zeros((len(cats_x), len(cats_y)), dtype=int)
        for i, j in zip(cx, cy):
            table[i, j] += 1

        active = [(i, j, table[i, j]) for i in range(table.shape[0]) for j in range(table.shape[1]) if table[i, j] > 0]
        lowers = np.asarray([[tx[i], ty[j]] for i, j, _ in active], dtype=float)
        uppers = np.asarray([[tx[i + 1], ty[j + 1]] for i, j, _ in active], dtype=float)
        weights = np.asarray([count for _, _, count in active], dtype=float)

        def objective(rho: float) -> float:
            cov = np.array([[1.0, rho], [rho, 1.0]])
            probs = multivariate_normal.cdf(
                uppers, mean=np.zeros(2), cov=cov, lower_limit=lowers,
                maxpts=10000, abseps=1e-5, releps=1e-5
            )
            probs = np.clip(np.asarray(probs, dtype=float), 1e-12, 1.0)
            return float(-np.sum(weights * np.log(probs)))

        result = minimize_scalar(objective, bounds=(-0.98, 0.98), method="bounded", options={"xatol": 5e-4, "maxiter": 60})
        if not result.success or not np.isfinite(result.x):
            return np.nan, "optimization_failed"
        rho = float(np.clip(result.x, -0.98, 0.98))
        status = "boundary" if abs(rho) >= 0.95 else "ok"
        return rho, status
    except Exception as exc:
        return np.nan, f"error:{type(exc).__name__}"


def polychoric_correlation(data: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    cols = list(data.columns)
    p = len(cols)
    mat = np.eye(p)
    pair_rows = []
    for i in range(p):
        for j in range(i + 1, p):
            rho, status = polychoric_pair(data.iloc[:, i].to_numpy(float), data.iloc[:, j].to_numpy(float))
            if not np.isfinite(rho):
                pair = data.iloc[:, [i, j]].corr(method="spearman").iloc[0, 1]
                rho = 0.0 if not np.isfinite(pair) else float(pair)
                fallback = True
            else:
                fallback = False
            mat[i, j] = mat[j, i] = rho
            pair_rows.append({
                "item_1": cols[i],
                "item_2": cols[j],
                "rho": rho,
                "status": status,
                "fallback_spearman": fallback,
                "pairwise_n": int(data.iloc[:, [i, j]].dropna().shape[0]),
            })
    fixed, near_diag = nearest_correlation(mat)
    pairs = pd.DataFrame(pair_rows)
    diagnostics = {
        **near_diag,
        "pair_count": int(len(pairs)),
        "fallback_count": int(pairs["fallback_spearman"].sum()) if len(pairs) else 0,
        "boundary_count": int((pairs["status"] == "boundary").sum()) if len(pairs) else 0,
        "fallback_fraction": float(pairs["fallback_spearman"].mean()) if len(pairs) else 0,
        "boundary_fraction": float((pairs["status"] == "boundary").mean()) if len(pairs) else 0,
    }
    return pd.DataFrame(fixed, index=cols, columns=cols), pairs, diagnostics


def kmo_from_corr(r: np.ndarray) -> tuple[np.ndarray, float]:
    r = np.asarray(r, dtype=float)
    inv = np.linalg.pinv(r)
    d = np.sqrt(np.clip(np.diag(inv), 1e-12, None))
    partial = -inv / np.outer(d, d)
    np.fill_diagonal(partial, 0.0)
    rr = r.copy()
    np.fill_diagonal(rr, 0.0)
    r2 = rr**2
    p2 = partial**2
    item = r2.sum(axis=0) / np.clip(r2.sum(axis=0) + p2.sum(axis=0), 1e-12, None)
    total = float(r2.sum() / np.clip(r2.sum() + p2.sum(), 1e-12, None))
    return item, total


def bartlett_from_corr(r: np.ndarray, n: int) -> tuple[float, int, float]:
    p = r.shape[0]
    sign, logdet = np.linalg.slogdet(r)
    if sign <= 0:
        return np.nan, p * (p - 1) // 2, np.nan
    stat = -(n - 1 - (2 * p + 5) / 6) * logdet
    df = p * (p - 1) // 2
    return float(stat), int(df), float(chi2.sf(stat, df))


def choose_correlation_method(
    item_diag: pd.DataFrame,
    n: int,
    poly_diag: dict[str, Any],
    cfg: Config,
) -> tuple[str, list[str]]:
    reasons = []
    ordinal = bool((item_diag["unique_categories"] <= 7).mean() >= 0.80)
    few_categories = bool(item_diag["unique_categories"].median() <= 5)
    skewed = bool((item_diag["skewness"].abs() > 1).mean() >= 0.20)
    unstable_poly = (
        poly_diag.get("fallback_fraction", 1) > 0.10
        or poly_diag.get("boundary_fraction", 1) > 0.10
        or poly_diag.get("maximum_absolute_adjustment", 1) > 0.10
    )
    sparse = bool((item_diag["sparse_category_fraction"] > 0.25).mean() >= 0.25)

    if ordinal and (few_categories or skewed):
        reasons.append("題目主要為少類別Likert序位資料，polychoric相關在概念上較符合量尺性質。")
        preferred = "polychoric"
    else:
        reasons.append("題目類別較多且分布較接近連續變項，Pearson相關可作為合理主分析。")
        preferred = "pearson"

    if unstable_poly:
        reasons.append("polychoric矩陣出現較多邊界估計、替代估計或正定化調整，顯示小樣本下可能不穩定。")
        preferred = "pearson"
    if n < 50 and sparse:
        reasons.append("樣本少於50且多個反應類別人數稀少；為避免polychoric估計過度波動，主分析採Pearson，polychoric保留為敏感度分析。")
        preferred = "pearson"
    if preferred == "polychoric" and n < 50:
        reasons.append("雖採polychoric為主分析，但樣本偏小，因此同時保留Pearson解並進行穩定性檢查。")
    return preferred, reasons


# -----------------------------------------------------------------------------
# Number of factors
# -----------------------------------------------------------------------------


def parallel_analysis_pearson(data: pd.DataFrame, iterations: int, seed: int) -> dict[str, Any]:
    rng = np.random.default_rng(seed)
    x = data.copy()
    x = x.fillna(x.median(numeric_only=True))
    obs = np.linalg.eigvalsh(x.corr().values)[::-1]
    random_eigs = np.empty((iterations, x.shape[1]))
    for b in range(iterations):
        z = rng.normal(size=x.shape)
        random_eigs[b] = np.linalg.eigvalsh(np.corrcoef(z, rowvar=False))[::-1]
    p95 = np.percentile(random_eigs, 95, axis=0)
    mean = np.mean(random_eigs, axis=0)
    k = int(np.sum(obs > p95))
    return {"observed": obs, "random_p95": p95, "random_mean": mean, "n_factors": max(k, 1)}


def simulate_independent_ordinal(data: pd.DataFrame, rng: np.random.Generator) -> pd.DataFrame:
    out = {}
    n = len(data)
    for col in data.columns:
        values = data[col].dropna().values
        categories, counts = np.unique(values, return_counts=True)
        probs = counts / counts.sum()
        out[col] = rng.choice(categories, size=n, replace=True, p=probs)
    return pd.DataFrame(out)


def parallel_analysis_polychoric(data: pd.DataFrame, observed_corr: pd.DataFrame, iterations: int, seed: int) -> dict[str, Any]:
    rng = np.random.default_rng(seed)
    obs = np.linalg.eigvalsh(observed_corr.values)[::-1]
    random_eigs = []
    failures = 0
    for b in range(iterations):
        sim = simulate_independent_ordinal(data, rng)
        try:
            r, _, _ = polychoric_correlation(sim)
            random_eigs.append(np.linalg.eigvalsh(r.values)[::-1])
        except Exception:
            failures += 1
    if not random_eigs:
        return {"observed": obs, "random_p95": np.full_like(obs, np.nan), "random_mean": np.full_like(obs, np.nan), "n_factors": np.nan, "failures": failures}
    random_eigs = np.asarray(random_eigs)
    p95 = np.percentile(random_eigs, 95, axis=0)
    mean = np.mean(random_eigs, axis=0)
    k = int(np.sum(obs > p95))
    return {"observed": obs, "random_p95": p95, "random_mean": mean, "n_factors": max(k, 1), "failures": failures}


def velicer_map(r: np.ndarray) -> pd.DataFrame:
    p = r.shape[0]
    vals, vecs = eigh(r)
    order = np.argsort(vals)[::-1]
    vals = vals[order]
    vecs = vecs[:, order]
    rows = []
    for m in range(0, max(1, p - 1)):
        if m == 0:
            residual = r.copy()
        else:
            reproduced = vecs[:, :m] @ np.diag(vals[:m]) @ vecs[:, :m].T
            residual = r - reproduced
        d = np.sqrt(np.clip(np.diag(residual), 1e-8, None))
        partial = residual / np.outer(d, d)
        np.fill_diagonal(partial, 0.0)
        off = offdiag_values(partial)
        rows.append({
            "components_partialled": m,
            "map_squared": float(np.mean(off**2)),
            "map_fourth": float(np.mean(off**4)),
        })
    return pd.DataFrame(rows)


def factor_count_evidence(
    data: pd.DataFrame,
    pearson_corr: pd.DataFrame,
    poly_corr: pd.DataFrame,
    cfg: Config,
) -> tuple[pd.DataFrame, dict[str, Any], dict[str, Any], pd.DataFrame, pd.DataFrame]:
    pa_p = parallel_analysis_pearson(data, cfg.pa_iterations_pearson, cfg.random_seed)
    poly_pa_iterations = max(10, int(cfg.pa_iterations_polychoric))
    pa_poly = parallel_analysis_polychoric(
        data,
        poly_corr,
        poly_pa_iterations,
        cfg.random_seed + 1,
    )
    pa_poly["iterations_requested"] = poly_pa_iterations
    map_p = velicer_map(pearson_corr.values)
    map_poly = velicer_map(poly_corr.values)

    rows = [
        {"method": "Parallel analysis (Pearson, 95th percentile)", "suggested_factors": pa_p["n_factors"], "role": "primary statistical evidence"},
        {"method": "Parallel analysis (polychoric, 95th percentile)", "suggested_factors": pa_poly["n_factors"], "role": "ordinal sensitivity evidence"},
        {"method": "Velicer MAP squared (Pearson)", "suggested_factors": int(map_p.loc[map_p["map_squared"].idxmin(), "components_partialled"]), "role": "sensitivity evidence"},
        {"method": "Revised MAP fourth power (Pearson)", "suggested_factors": int(map_p.loc[map_p["map_fourth"].idxmin(), "components_partialled"]), "role": "sensitivity evidence"},
        {"method": "Velicer MAP squared (polychoric)", "suggested_factors": int(map_poly.loc[map_poly["map_squared"].idxmin(), "components_partialled"]), "role": "ordinal sensitivity evidence"},
    ]
    if cfg.expected_factors:
        rows.append({"method": "Planned theoretical dimensions", "suggested_factors": int(cfg.expected_factors), "role": "theoretical evidence"})
    evidence = pd.DataFrame(rows)
    evidence.loc[evidence["suggested_factors"] < 1, "suggested_factors"] = 1
    return evidence, pa_p, pa_poly, map_p, map_poly


# -----------------------------------------------------------------------------
# Factor models, alignment, and model comparison
# -----------------------------------------------------------------------------


def planned_factor_info(codebook: pd.DataFrame, items: list[str]) -> tuple[list[str], dict[str, str]]:
    if codebook.empty or "expected_factor" not in codebook.columns:
        return [], {}
    temp = codebook.set_index("item").reindex(items)
    mapping = {
        item: str(temp.loc[item, "expected_factor"]).strip()
        for item in items
        if item in temp.index and pd.notna(temp.loc[item, "expected_factor"]) and str(temp.loc[item, "expected_factor"]).strip()
    }
    levels = []
    for item in items:
        f = mapping.get(item)
        if f and f not in levels:
            levels.append(f)
    return levels, mapping


def orient_and_align(
    loadings: np.ndarray,
    phi: np.ndarray,
    items: list[str],
    expected_levels: list[str],
    expected_map: dict[str, str],
) -> tuple[np.ndarray, np.ndarray, list[str], dict[str, Any]]:
    k = loadings.shape[1]
    alignment: dict[str, Any] = {"used_theoretical_alignment": False}
    labels = [f"F{i+1}" for i in range(k)]
    order = list(range(k))

    if expected_levels and k == len(expected_levels):
        score = np.zeros((len(expected_levels), k))
        for i, level in enumerate(expected_levels):
            idx = [j for j, item in enumerate(items) if expected_map.get(item) == level]
            if idx:
                score[i, :] = np.mean(np.abs(loadings[idx, :]), axis=0)
        rows, cols = linear_sum_assignment(-score)
        assignment = {expected_levels[r]: int(c) for r, c in zip(rows, cols)}
        order = [assignment[level] for level in expected_levels]
        labels = expected_levels.copy()
        alignment.update({"used_theoretical_alignment": True, "alignment_scores": score.tolist(), "assignment": assignment})
    else:
        # Order by sum of squared loadings.
        order = list(np.argsort(np.sum(loadings**2, axis=0))[::-1])

    loadings = loadings[:, order]
    phi = reorder_square(phi, order)

    # Orient each factor so that expected or strongest salient loadings are positive.
    signs = np.ones(k)
    for j in range(k):
        if expected_levels and labels[j] in expected_levels:
            idx = [i for i, item in enumerate(items) if expected_map.get(item) == labels[j]]
            vals = loadings[idx, j] if idx else loadings[:, j]
        else:
            vals = loadings[:, j]
        salient = vals[np.abs(vals) >= 0.30]
        target = np.nanmedian(salient) if len(salient) else vals[np.argmax(np.abs(vals))]
        if target < 0:
            signs[j] = -1
    loadings = loadings * signs
    phi = phi * np.outer(signs, signs)
    return loadings, phi, labels, alignment


def fit_factor_model(
    corr: pd.DataFrame,
    n: int,
    k: int,
    extraction: str,
    rotation: str,
    items: list[str],
    expected_levels: list[str],
    expected_map: dict[str, str],
) -> dict[str, Any]:
    if k < 1 or k >= len(items):
        raise ValueError("因素數必須介於1與題目數減1之間")
    method = "pa" if extraction in {"pa", "principal_axis"} else "ml"
    model = Factor(corr=corr.values, n_factor=k, method=method, smc=True, nobs=n, endog_names=items)
    result = model.fit(maxiter=1000, tol=1e-7)
    if k > 1 and rotation != "none":
        result.rotate(rotation)
    loadings = np.asarray(result.loadings, dtype=float)
    if rotation in {"oblimin", "promax"} and k > 1:
        t = np.asarray(result.rotation_matrix, dtype=float)
        phi = t.T @ t
        d = np.sqrt(np.clip(np.diag(phi), 1e-12, None))
        phi = phi / np.outer(d, d)
    else:
        phi = np.eye(k)
    loadings, phi, labels, alignment = orient_and_align(loadings, phi, items, expected_levels, expected_map)
    structure = loadings @ phi
    communalities = np.diag(loadings @ phi @ loadings.T)
    uniqueness = 1 - communalities
    implied = loadings @ phi @ loadings.T + np.diag(uniqueness)
    residual = corr.values - implied
    rmsr = float(np.sqrt(np.mean(offdiag_values(residual) ** 2)))
    complexity = (np.sum(loadings**2, axis=1) ** 2) / np.clip(np.sum(loadings**4, axis=1), 1e-12, None)

    abs_load = np.abs(loadings)
    primary_idx = np.argmax(abs_load, axis=1)
    sorted_load = np.sort(abs_load, axis=1)[:, ::-1]
    primary = sorted_load[:, 0]
    secondary = sorted_load[:, 1] if k > 1 else np.zeros(len(items))
    gap = primary - secondary
    assigned = [labels[j] for j in primary_idx]
    counts = Counter(assigned)
    t = Thresholds()
    item_table = pd.DataFrame({
        "item": items,
        "assigned_factor": assigned,
        "primary_loading": primary,
        "secondary_loading": secondary,
        "loading_gap": gap,
        "communality": communalities,
        "uniqueness": uniqueness,
        "complexity": complexity,
    })
    for j, label in enumerate(labels):
        item_table[f"loading_{label}"] = loadings[:, j]
        item_table[f"structure_{label}"] = structure[:, j]

    mismatch = 0
    if expected_map and k == len(expected_levels):
        item_table["expected_factor"] = [expected_map.get(x, "") for x in items]
        item_table["matches_expected"] = item_table["expected_factor"] == item_table["assigned_factor"]
        mismatch = int((~item_table["matches_expected"] & (item_table["expected_factor"] != "")).sum())

    metrics = {
        "k": k,
        "extraction": extraction,
        "rotation": rotation,
        "rmsr": rmsr,
        "mean_primary_loading": float(np.mean(primary)),
        "median_primary_loading": float(np.median(primary)),
        "low_loading_count": int(np.sum(primary < t.loading_minimum)),
        "weak_loading_count": int(np.sum(primary < t.loading_preferred)),
        "cross_loading_count": int(np.sum((secondary >= t.cross_loading) & (gap < t.loading_gap))),
        "low_communality_count": int(np.sum(communalities < t.communality_severe)),
        "factor_below_three_items": int(sum(v < 3 for v in counts.values()) + max(0, k - len(counts))),
        "heywood_count": int(np.sum((uniqueness <= 0.005) | (uniqueness > 1.0))),
        "mean_complexity": float(np.mean(complexity)),
        "theoretical_mismatch_count": mismatch,
        "maximum_factor_correlation": float(np.max(np.abs(offdiag_values(phi)))) if k > 1 else 0.0,
        "factor_item_counts": dict(counts),
    }
    metrics["penalty"] = model_penalty(metrics)
    return {
        "result": result,
        "loadings": pd.DataFrame(loadings, index=items, columns=labels),
        "structure": pd.DataFrame(structure, index=items, columns=labels),
        "phi": pd.DataFrame(phi, index=labels, columns=labels),
        "communalities": pd.Series(communalities, index=items, name="communality"),
        "uniqueness": pd.Series(uniqueness, index=items, name="uniqueness"),
        "item_table": item_table,
        "metrics": metrics,
        "labels": labels,
        "alignment": alignment,
        "residual": pd.DataFrame(residual, index=items, columns=items),
    }


def model_penalty(m: dict[str, Any]) -> float:
    return float(
        25 * m.get("heywood_count", 0)
        + 12 * m.get("factor_below_three_items", 0)
        + 5 * m.get("low_loading_count", 0)
        + 4 * m.get("cross_loading_count", 0)
        + 3 * m.get("low_communality_count", 0)
        + 1.5 * m.get("weak_loading_count", 0)
        + 1.5 * m.get("theoretical_mismatch_count", 0)
        + 80 * m.get("rmsr", 0)
        + 2 * max(0.0, m.get("mean_complexity", 1) - 1)
        - 8 * m.get("mean_primary_loading", 0)
    )


def candidate_factor_numbers(evidence: pd.DataFrame, p: int, expected: Optional[int]) -> list[int]:
    """Return evidence-supported candidate factor counts.

    A one-factor solution is included only when suggested by evidence or when no
    multidimensional candidate exists. This prevents one factor from winning merely
    because it cannot display cross-loadings by definition.
    """
    vals = [int(v) for v in evidence["suggested_factors"].dropna().tolist() if int(v) >= 1]
    if expected:
        vals.append(int(expected))
    if vals:
        mode = Counter(vals).most_common(1)[0][0]
        vals += [mode - 1, mode + 1]
    if not vals:
        vals = [1]
    max_k = min(max(1, p // 2), 10, p - 2)
    return sorted({v for v in vals if 1 <= v <= max_k})


def compare_models(
    corr: pd.DataFrame,
    n: int,
    items: list[str],
    evidence: pd.DataFrame,
    cfg: Config,
    codebook: pd.DataFrame,
    extraction: str = "pa",
) -> tuple[pd.DataFrame, dict[tuple[int, str], dict[str, Any]]]:
    expected_levels, expected_map = planned_factor_info(codebook, items)
    expected_k = cfg.expected_factors or (len(expected_levels) if expected_levels else None)
    ks = candidate_factor_numbers(evidence, len(items), expected_k)
    models = {}
    rows = []
    evidence_counts = Counter(int(v) for v in evidence["suggested_factors"].dropna())
    for k in ks:
        for rotation in (["none"] if k == 1 else ["oblimin", "varimax"]):
            try:
                fit = fit_factor_model(corr, n, k, extraction, rotation, items, expected_levels, expected_map)
                vote_bonus = -2 * evidence_counts.get(k, 0)
                theory_bonus = -3 if expected_k and k == expected_k else 0
                fit["metrics"]["decision_score"] = fit["metrics"]["penalty"] + vote_bonus + theory_bonus
                rows.append({**fit["metrics"], "evidence_votes": evidence_counts.get(k, 0), "theoretical_k": expected_k, "status": "ok"})
                models[(k, rotation)] = fit
            except Exception as exc:
                rows.append({"k": k, "extraction": extraction, "rotation": rotation, "status": f"failed: {type(exc).__name__}: {exc}", "penalty": np.inf, "decision_score": np.inf})
    table = pd.DataFrame(rows).sort_values(["decision_score", "penalty"], na_position="last")
    return table, models


def choose_model(
    comparison: pd.DataFrame,
    models: dict[tuple[int, str], dict[str, Any]],
    cfg: Config,
    codebook: pd.DataFrame,
    items: list[str],
    n: int,
    evidence: pd.DataFrame,
) -> tuple[dict[str, Any], list[str]]:
    ok = comparison[comparison["status"] == "ok"].copy()
    if ok.empty:
        raise RuntimeError("所有候選因素模型都估計失敗。")
    expected_levels, _ = planned_factor_info(codebook, items)
    expected_k = cfg.expected_factors or (len(expected_levels) if expected_levels else None)
    reasons: list[str] = []

    # Exclude unsupported one-factor solutions from the primary choice when a
    # multidimensional questionnaire was planned and neither parallel analysis
    # supports one factor. It remains available in the detailed workbook if fitted.
    pa_rows = evidence[evidence["method"].str.contains("Parallel analysis", case=False, na=False)]
    pa_supported = set(int(v) for v in pa_rows["suggested_factors"].dropna())
    eligible = ok.copy()
    if expected_k and expected_k > 1 and 1 not in pa_supported:
        eligible = eligible[eligible["k"] != 1]
        reasons.append("單因素解未獲平行分析支持，因此只保留為敏感度結果，不列入主要建議候選。")
    if eligible.empty:
        eligible = ok.copy()

    best_row = eligible.iloc[0]
    chosen_row = best_row

    if expected_k and cfg.prioritize_planned_dimensions:
        theory_rows = eligible[eligible["k"] == expected_k]
        if not theory_rows.empty:
            theory_best = theory_rows.iloc[0]
            problem_count = (
                safe_float(theory_best.get("low_loading_count"), 0)
                + safe_float(theory_best.get("cross_loading_count"), 0)
                + safe_float(theory_best.get("low_communality_count"), 0)
            )
            problem_fraction = problem_count / max(len(items), 1)
            strong_counterevidence = (
                safe_float(theory_best.get("factor_below_three_items"), 0) > 0
                or safe_float(theory_best.get("heywood_count"), 0) > 0
                or safe_float(theory_best.get("maximum_factor_correlation"), 0) >= cfg.planned_factor_correlation_critical
                or safe_float(theory_best.get("rmsr"), 0) >= cfg.planned_rmsr_critical
                or problem_fraction >= cfg.planned_problem_fraction_critical
            )
            if n < cfg.preserve_planned_below_n:
                chosen_row = theory_best
                reasons.append(
                    f"樣本數N={n}低於{cfg.preserve_planned_below_n}；單次EFA不足以永久推翻原定{expected_k}個section，因此以原定結構作暫時工作方案，其他因素解作敏感度分析。"
                )
            elif not strong_counterevidence:
                chosen_row = theory_best
                reasons.append(
                    f"原定{expected_k}因素解未出現預先設定的強烈反證，故優先保留理論結構。"
                )
            else:
                alternative_supported = int(best_row["k"]) in pa_supported
                material_improvement = safe_float(theory_best.get("decision_score"), np.inf) - safe_float(best_row.get("decision_score"), np.inf) >= 5
                if alternative_supported and material_improvement:
                    chosen_row = best_row
                    reasons.append(
                        f"原定{expected_k}因素解出現強烈結構警示，且{int(best_row['k'])}因素解獲平行分析支持並有實質改善，因此建議以{int(best_row['k'])}因素作主要統計方案；原定方案仍完整保留。"
                    )
                else:
                    chosen_row = theory_best
                    reasons.append(
                        f"原定{expected_k}因素解雖有警示，但替代因素數的證據或改善幅度不足，故暫時保留原定方案並標示限制。"
                    )

    chosen = models[(int(chosen_row["k"]), str(chosen_row["rotation"]))]
    if chosen_row["rotation"] == "oblimin":
        reasons.append(f"斜交解的最大因素相關為{fmt(chosen['metrics']['maximum_factor_correlation'])}；因素可能相關，因此採斜交轉軸（oblique rotation）。")
    elif chosen_row["rotation"] == "varimax":
        reasons.append("斜交解顯示因素相關很低，且正交解較簡潔，因此採Varimax正交轉軸（orthogonal rotation）。")
    else:
        reasons.append("單因素模型不需要轉軸。")
    return chosen, reasons


# -----------------------------------------------------------------------------
# Item diagnostics and iterative deletion
# -----------------------------------------------------------------------------


def redundancy_flags(corr: pd.DataFrame, threshold: float) -> dict[str, list[str]]:
    flags: dict[str, list[str]] = defaultdict(list)
    for i, a in enumerate(corr.columns):
        for j in range(i + 1, len(corr.columns)):
            b = corr.columns[j]
            if abs(corr.iloc[i, j]) >= threshold:
                flags[a].append(b)
                flags[b].append(a)
    return flags


def diagnose_items(
    fit: dict[str, Any],
    corr: pd.DataFrame,
    msa: pd.Series,
    cfg: Config,
    expected_map: dict[str, str],
) -> pd.DataFrame:
    t = cfg.thresholds
    df = fit["item_table"].copy().set_index("item")
    df["msa"] = msa.reindex(df.index)
    redundancy = redundancy_flags(corr, t.redundancy_correlation)
    reasons = []
    severity = []
    scores = []
    for item, row in df.iterrows():
        r = []
        severe = False
        score = 0.0
        if row["primary_loading"] < t.loading_minimum:
            r.append(f"主要負荷{row['primary_loading']:.3f}< {t.loading_minimum:.2f}")
            severe = True
            score += 5
        elif row["primary_loading"] < t.loading_preferred:
            r.append(f"主要負荷僅{row['primary_loading']:.3f}")
            score += 1.5
        if row["secondary_loading"] >= t.cross_loading and row["loading_gap"] < t.loading_gap:
            r.append(f"交叉負荷：次高負荷{row['secondary_loading']:.3f}，差距{row['loading_gap']:.3f}")
            severe = True
            score += 5
        elif row["secondary_loading"] >= 0.25 and row["loading_gap"] < 0.25:
            r.append("有輕度交叉負荷")
            score += 1
        if row["communality"] < t.communality_severe:
            r.append(f"共同性{row['communality']:.3f}< {t.communality_severe:.2f}")
            severe = True
            score += 4
        elif row["communality"] < t.communality_warning:
            r.append(f"共同性偏低({row['communality']:.3f})")
            score += 1
        if np.isfinite(row["msa"]) and row["msa"] < t.msa_minimum:
            r.append(f"MSA={row['msa']:.3f}< {t.msa_minimum:.2f}")
            severe = True
            score += 3
        if redundancy.get(item):
            r.append(f"與{', '.join(redundancy[item])}高度相關")
            score += 1.5
        expected = expected_map.get(item)
        if expected and fit["metrics"]["k"] == len(set(expected_map.values())) and row["assigned_factor"] != expected:
            r.append(f"主要落在{row['assigned_factor']}而非預定的{expected}")
            score += 2
        reasons.append("；".join(r) if r else "無明顯問題")
        severity.append("severe" if severe else ("warning" if r else "ok"))
        scores.append(score)
    df["diagnostic_level"] = severity
    df["diagnostic_score"] = scores
    df["diagnostic_reasons"] = reasons
    return df.reset_index()


def eligible_to_delete(
    item: str,
    current_items: list[str],
    diagnostics: pd.DataFrame,
    expected_map: dict[str, str],
    minimum_items: int,
    protected_items: Optional[set[str]] = None,
) -> bool:
    if protected_items and item in protected_items:
        return False
    if expected_map.get(item):
        f = expected_map[item]
        remaining = sum(1 for i in current_items if i != item and expected_map.get(i) == f)
        return remaining >= minimum_items
    assigned = diagnostics.set_index("item").loc[item, "assigned_factor"]
    remaining = sum(1 for i in current_items if i != item and diagnostics.set_index("item").loc[i, "assigned_factor"] == assigned)
    return remaining >= minimum_items


def quick_fit_penalty_after_deletion(
    data: pd.DataFrame,
    items: list[str],
    delete_item: str,
    method: str,
    k: int,
    rotation: str,
    codebook: pd.DataFrame,
) -> float:
    subset = [x for x in items if x != delete_item]
    if k >= len(subset):
        return np.inf
    try:
        if method == "polychoric":
            corr, _, _ = polychoric_correlation(data[subset])
        else:
            corr, _ = pearson_correlation(data[subset])
        levels, mapping = planned_factor_info(codebook, subset)
        fit = fit_factor_model(corr, len(data), k, "pa", rotation, subset, levels, mapping)
        return fit["metrics"]["penalty"]
    except Exception:
        return np.inf


def iterative_item_refinement(
    data: pd.DataFrame,
    initial_items: list[str],
    codebook: pd.DataFrame,
    cfg: Config,
    primary_method: str,
    initial_k: int,
    initial_rotation: str,
) -> dict[str, Any]:
    current_items = initial_items.copy()
    deleted = []
    iterations = []
    alternatives = []
    max_delete = max(0, int(math.floor(len(initial_items) * cfg.max_delete_fraction)))
    if len(data) < cfg.disable_auto_deletion_below_n:
        max_delete = 0
    expected_levels_all, expected_map_all = planned_factor_info(codebook, initial_items)
    protected_items: set[str] = set()
    if not codebook.empty and "protect" in codebook.columns:
        protected_items = set(
            codebook.loc[codebook["protect"].map(boolish), "item"].astype(str).str.strip()
        )
    k = initial_k
    rotation = initial_rotation
    final_fit = None
    final_corr = None
    final_diag = None

    for iteration in range(cfg.max_delete_iterations + 1):
        if primary_method == "polychoric":
            corr, _, corr_diag = polychoric_correlation(data[current_items])
        else:
            corr, corr_diag = pearson_correlation(data[current_items])
        msa_vals, kmo_total = kmo_from_corr(corr.values)
        msa = pd.Series(msa_vals, index=current_items)
        levels, mapping = planned_factor_info(codebook, current_items)
        if k >= len(current_items):
            k = max(1, len(current_items) - 2)
        try:
            # Re-check both rotations every iteration.
            fits = {}
            for rot in (["none"] if k == 1 else ["oblimin", "varimax"]):
                try:
                    fits[rot] = fit_factor_model(corr, len(data), k, "pa", rot, current_items, levels, mapping)
                except Exception:
                    pass
            if not fits:
                raise RuntimeError("本迭代所有轉軸都失敗")
            if k == 1:
                rotation = "none"
            else:
                ob = fits.get("oblimin")
                va = fits.get("varimax")
                if ob is not None and ob["metrics"]["maximum_factor_correlation"] >= cfg.thresholds.factor_correlation_oblique:
                    rotation = "oblimin"
                elif va is not None and (ob is None or va["metrics"]["penalty"] + 1 < ob["metrics"]["penalty"]):
                    rotation = "varimax"
                else:
                    rotation = "oblimin" if ob is not None else "varimax"
            fit = fits[rotation]
        except Exception as exc:
            iterations.append({"iteration": iteration, "status": f"fit_failed: {type(exc).__name__}: {exc}", "n_items": len(current_items)})
            break

        diagnostics = diagnose_items(fit, corr, msa, cfg, mapping)
        severe = diagnostics[diagnostics["diagnostic_level"] == "severe"].copy()
        iterations.append({
            "iteration": iteration,
            "status": "estimated",
            "n_items": len(current_items),
            "k": k,
            "rotation": rotation,
            "kmo": kmo_total,
            **fit["metrics"],
            "severe_item_count": int(len(severe)),
            "current_items": ", ".join(current_items),
        })
        final_fit, final_corr, final_diag = fit, corr, diagnostics

        if severe.empty:
            break
        if max_delete == 0 and len(data) < cfg.disable_auto_deletion_below_n:
            iterations[-1]["stop_reason"] = f"樣本數N={len(data)}低於{cfg.disable_auto_deletion_below_n}，不自動永久刪題；問題題目僅列為修訂候選"
            break
        if len(deleted) >= max_delete:
            iterations[-1]["stop_reason"] = "已達最大刪題比例"
            break

        severe = severe.sort_values(["diagnostic_score", "primary_loading"], ascending=[False, True])
        candidates = [
            item for item in severe["item"].tolist()
            if eligible_to_delete(
                item, current_items, diagnostics, mapping or expected_map_all,
                cfg.minimum_items_per_factor, protected_items
            )
        ]
        if not candidates:
            iterations[-1]["stop_reason"] = "剩餘問題題目若刪除會使某因素少於最低題數"
            break

        top = candidates[:2]
        penalties = {}
        for item in top:
            penalties[item] = quick_fit_penalty_after_deletion(data, current_items, item, primary_method, k, rotation, codebook)
        chosen_item = min(penalties, key=penalties.get)
        if len(top) == 2:
            alternatives.append({
                "iteration": iteration,
                "candidate_1": top[0],
                "penalty_after_deleting_1": penalties[top[0]],
                "candidate_2": top[1],
                "penalty_after_deleting_2": penalties[top[1]],
                "chosen": chosen_item,
            })

        row = diagnostics.set_index("item").loc[chosen_item]
        deleted.append({
            "iteration": iteration + 1,
            "item": chosen_item,
            "reason": row["diagnostic_reasons"],
            "primary_loading": row["primary_loading"],
            "secondary_loading": row["secondary_loading"],
            "loading_gap": row["loading_gap"],
            "communality": row["communality"],
            "msa": row["msa"],
            "assigned_factor": row["assigned_factor"],
            "expected_factor": mapping.get(chosen_item, expected_map_all.get(chosen_item, "")),
            "decision_note": "一次只刪除一題；若有兩個相近候選，已比較刪除後模型懲罰值。",
        })
        current_items.remove(chosen_item)

        # Keep the selected factor count fixed during item-refinement.
        # Changing k after each deletion would confound item decisions with a new
        # dimensionality decision and could pull an alternative solution back to
        # the planned number of factors. Factor-count sensitivity is handled in
        # the separate candidate-model comparison instead.
        k = min(k, max(1, len(current_items) - 2))

    return {
        "final_items": current_items,
        "deleted_items": pd.DataFrame(deleted),
        "iteration_log": pd.DataFrame(iterations),
        "alternative_deletions": pd.DataFrame(alternatives),
        "final_fit": final_fit,
        "final_corr": final_corr,
        "final_diagnostics": final_diag,
    }


# -----------------------------------------------------------------------------
# Reliability and validity
# -----------------------------------------------------------------------------


def alpha_from_corr(r: np.ndarray) -> float:
    k = r.shape[0]
    if k < 2:
        return np.nan
    total = np.sum(r)
    return float(k / (k - 1) * (1 - k / total)) if total > 0 else np.nan


def cronbach_alpha(data: pd.DataFrame) -> float:
    x = data.dropna(how="all").copy()
    if x.shape[1] < 2:
        return np.nan
    x = x.apply(lambda c: c.fillna(c.median()))
    item_var = x.var(axis=0, ddof=1).sum()
    total_var = x.sum(axis=1).var(ddof=1)
    k = x.shape[1]
    return float(k / (k - 1) * (1 - item_var / total_var)) if total_var > 0 else np.nan


def omega_one_factor(corr: pd.DataFrame, n: int) -> tuple[float, np.ndarray]:
    if corr.shape[0] < 2:
        return np.nan, np.array([])
    try:
        fit = Factor(corr=corr.values, n_factor=1, method="pa", smc=True, nobs=n).fit()
        lam = np.asarray(fit.loadings[:, 0], dtype=float)
        if np.nanmedian(lam) < 0:
            lam = -lam
        uniq = np.clip(1 - lam**2, 0, None)
        omega = (np.sum(lam) ** 2) / ((np.sum(lam) ** 2) + np.sum(uniq))
        return float(omega), lam
    except Exception:
        return np.nan, np.full(corr.shape[0], np.nan)


def bootstrap_stat_ci(data: pd.DataFrame, func, iterations: int, seed: int) -> tuple[float, float, int]:
    rng = np.random.default_rng(seed)
    values = []
    n = len(data)
    for _ in range(iterations):
        sample = data.iloc[rng.integers(0, n, n)]
        try:
            value = func(sample)
            if np.isfinite(value):
                values.append(value)
        except Exception:
            pass
    if len(values) < max(20, iterations * 0.25):
        return np.nan, np.nan, len(values)
    return float(np.percentile(values, 2.5)), float(np.percentile(values, 97.5)), len(values)


def htmt_matrix(data: pd.DataFrame, groups: dict[str, list[str]]) -> pd.DataFrame:
    labels = list(groups)
    r = data[list({i for g in groups.values() for i in g})].corr().abs()
    out = pd.DataFrame(np.eye(len(labels)), index=labels, columns=labels)
    for a_idx, a in enumerate(labels):
        for b_idx in range(a_idx + 1, len(labels)):
            b = labels[b_idx]
            hetero = [r.loc[i, j] for i in groups[a] for j in groups[b] if i in r.index and j in r.columns]
            mono_a = [r.loc[groups[a][i], groups[a][j]] for i in range(len(groups[a])) for j in range(i + 1, len(groups[a]))]
            mono_b = [r.loc[groups[b][i], groups[b][j]] for i in range(len(groups[b])) for j in range(i + 1, len(groups[b]))]
            denom = math.sqrt(max(np.mean(mono_a), 1e-12) * max(np.mean(mono_b), 1e-12)) if mono_a and mono_b else np.nan
            val = float(np.mean(hetero) / denom) if hetero and np.isfinite(denom) else np.nan
            out.loc[a, b] = out.loc[b, a] = val
    return out


def reliability_validity(
    data: pd.DataFrame,
    final_fit: dict[str, Any],
    final_items: list[str],
    poly_corr_final: pd.DataFrame,
    cfg: Config,
) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    assignments = final_fit["item_table"].set_index("item")["assigned_factor"].to_dict()
    groups = defaultdict(list)
    for item in final_items:
        groups[assignments[item]].append(item)

    rel_rows = []
    valid_rows = []
    for idx, (factor, items) in enumerate(groups.items()):
        sub = data[items]
        pearson_r = sub.corr()
        poly_r = poly_corr_final.loc[items, items]
        alpha = cronbach_alpha(sub)
        ordinal_alpha = alpha_from_corr(poly_r.values)
        omega, omega_load = omega_one_factor(pearson_r, len(sub))
        ordinal_omega, ord_load = omega_one_factor(poly_r, len(sub))
        alpha_l, alpha_u, alpha_b = bootstrap_stat_ci(sub, cronbach_alpha, cfg.bootstrap_iterations, cfg.random_seed + 100 + idx)
        omega_l, omega_u, omega_b = bootstrap_stat_ci(
            sub,
            lambda d: omega_one_factor(d.corr(), len(d))[0],
            cfg.bootstrap_iterations,
            cfg.random_seed + 200 + idx,
        )
        rel_rows.append({
            "factor": factor,
            "items": ", ".join(items),
            "n_items": len(items),
            "cronbach_alpha": alpha,
            "alpha_ci_lower": alpha_l,
            "alpha_ci_upper": alpha_u,
            "alpha_bootstrap_success": alpha_b,
            "ordinal_alpha": ordinal_alpha,
            "mcdonald_omega": omega,
            "omega_ci_lower": omega_l,
            "omega_ci_upper": omega_u,
            "omega_bootstrap_success": omega_b,
            "ordinal_omega": ordinal_omega,
        })

        final_loading_cols = [c for c in final_fit["item_table"].columns if c.startswith("loading_")]
        col = f"loading_{factor}"
        if col in final_fit["item_table"].columns:
            lam = final_fit["item_table"].set_index("item").loc[items, col].abs().values
            ave = float(np.mean(lam**2))
            cr = float((np.sum(lam) ** 2) / ((np.sum(lam) ** 2) + np.sum(1 - lam**2)))
        else:
            ave = np.nan
            cr = np.nan
        valid_rows.append({
            "factor": factor,
            "average_variance_extracted_AVE": ave,
            "composite_reliability_CR": cr,
            "mean_primary_loading": float(final_fit["item_table"].set_index("item").loc[items, "primary_loading"].mean()),
            "minimum_primary_loading": float(final_fit["item_table"].set_index("item").loc[items, "primary_loading"].min()),
        })

    htmt = htmt_matrix(data[final_items], groups) if len(groups) > 1 else pd.DataFrame()
    return pd.DataFrame(rel_rows), pd.DataFrame(valid_rows), htmt


# -----------------------------------------------------------------------------
# Stability analyses
# -----------------------------------------------------------------------------


def align_bootstrap_loadings(reference: np.ndarray, candidate: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    ref = reference.copy()
    cand = candidate.copy()
    k = ref.shape[1]
    congr = np.zeros((k, k))
    for i in range(k):
        for j in range(k):
            denom = np.sqrt(np.sum(ref[:, i] ** 2) * np.sum(cand[:, j] ** 2))
            congr[i, j] = abs(float(np.sum(ref[:, i] * cand[:, j]) / denom)) if denom > 0 else 0
    rows, cols = linear_sum_assignment(-congr)
    order = [cols[np.where(rows == i)[0][0]] for i in range(k)]
    cand = cand[:, order]
    for j in range(k):
        if np.sum(ref[:, j] * cand[:, j]) < 0:
            cand[:, j] *= -1
    return cand, congr[:, order]


def bootstrap_stability(
    data: pd.DataFrame,
    final_fit: dict[str, Any],
    final_items: list[str],
    method: str,
    cfg: Config,
) -> pd.DataFrame:
    ref = final_fit["loadings"].loc[final_items].values
    k = ref.shape[1]
    rotation = final_fit["metrics"]["rotation"]
    levels, mapping = [], {}
    rng = np.random.default_rng(cfg.random_seed + 500)
    b_total = (
        min(cfg.bootstrap_polychoric_iterations, max(10, int(300 / max(len(final_items), 1))))
        if method == "polychoric" else cfg.bootstrap_iterations
    )
    arrays = []
    assignments = []
    success = 0
    for _ in range(b_total):
        sample = data.iloc[rng.integers(0, len(data), len(data))][final_items]
        try:
            if method == "polychoric":
                r, _, _ = polychoric_correlation(sample)
            else:
                r, _ = pearson_correlation(sample)
            fit = fit_factor_model(r, len(sample), k, "pa", rotation, final_items, levels, mapping)
            aligned, _ = align_bootstrap_loadings(ref, fit["loadings"].values)
            arrays.append(aligned)
            assignments.append(np.argmax(np.abs(aligned), axis=1))
            success += 1
        except Exception:
            pass
    if not arrays:
        return pd.DataFrame({"item": final_items, "bootstrap_success": 0})
    arr = np.stack(arrays)
    ass = np.stack(assignments)
    ref_ass = np.argmax(np.abs(ref), axis=1)
    rows = []
    for i, item in enumerate(final_items):
        j = ref_ass[i]
        vals = arr[:, i, j]
        rows.append({
            "item": item,
            "reference_factor": final_fit["loadings"].columns[j],
            "reference_loading": ref[i, j],
            "bootstrap_loading_median": float(np.median(vals)),
            "bootstrap_loading_ci_lower": float(np.percentile(vals, 2.5)),
            "bootstrap_loading_ci_upper": float(np.percentile(vals, 97.5)),
            "factor_assignment_stability": float(np.mean(ass[:, i] == j)),
            "bootstrap_success": success,
            "bootstrap_requested": b_total,
        })
    return pd.DataFrame(rows)


# -----------------------------------------------------------------------------
# Confirmatory feasibility and R script generation
# -----------------------------------------------------------------------------


def confirmatory_feasibility(n: int, p: int, k: int, min_items: int) -> dict[str, Any]:
    rough_parameters = p + p + k * (k - 1) / 2 + max(0, p - k)
    n_per_parameter = n / max(rough_parameters, 1)
    cfa_feasible = bool(n >= 200 and n >= 10 * p and min_items >= 3)
    diagnostic_cfa = bool(n >= 100 and n >= 5 * p and min_items >= 3)
    esem_feasible = bool(n >= 300 and n >= 12 * p and min_items >= 3)
    if cfa_feasible:
        decision = "可使用獨立樣本或分割樣本進行CFA；若本資料已用於刪題，同一樣本CFA仍只能視為內部診斷。"
    elif diagnostic_cfa:
        decision = "可做同樣本CFA作診斷，但不應稱為獨立驗證；不建議再切分樣本。"
    else:
        decision = "樣本不足以支持穩定CFA或ESEM；主結果應停留在探索性證據並強調不確定性。"
    return {
        "n": n,
        "p_final": p,
        "k_final": k,
        "rough_free_parameters": rough_parameters,
        "n_per_parameter": n_per_parameter,
        "cfa_independent_feasible": cfa_feasible,
        "cfa_diagnostic_only": diagnostic_cfa and not cfa_feasible,
        "esem_feasible": esem_feasible,
        "decision": decision,
    }


def r_quote(s: str) -> str:
    return '"' + s.replace('\\', '\\\\').replace('"', '\\"') + '"'


def generate_lavaan_scripts(
    output_dir: Path,
    input_excel: str,
    data_sheet: str,
    final_fit: dict[str, Any],
    final_items: list[str],
    feasibility: dict[str, Any],
) -> None:
    assignment = final_fit["item_table"].set_index("item")["assigned_factor"].to_dict()
    original_groups = defaultdict(list)
    for item in final_items:
        original_groups[assignment[item]].append(item)
    factor_name_map = {name: f"F{i+1}" for i, name in enumerate(original_groups)}
    groups = {factor_name_map[name]: items for name, items in original_groups.items()}
    mapping_comments = "\n".join([f"# {factor_name_map[name]} = {name}" for name in original_groups])
    model_lines = [f"{factor} =~ " + " + ".join(items) for factor, items in groups.items()]
    model_text = "\n".join(model_lines)
    ordered = ", ".join(r_quote(x) for x in final_items)
    factor_names = list(groups)
    efa_left = " +\n    ".join([f'efa("efa1")*{f}' for f in factor_names])
    all_items = " + ".join(final_items)

    cfa_script = f'''# Auto-generated by questionnaire_quality_pipeline.py
# This script uses WLSMV for ordinal Likert items.
{mapping_comments}
library(readxl)
library(lavaan)
library(semTools)

input_file <- {r_quote(str(Path(input_excel).resolve()))}
data <- read_excel(input_file, sheet = {r_quote(data_sheet)})
items <- c({ordered})
data[items] <- lapply(data[items], ordered)

model <- '
{model_text}
'

fit <- cfa(model, data = data, ordered = items, estimator = "WLSMV", std.lv = TRUE)
summary(fit, fit.measures = TRUE, standardized = TRUE)
write.csv(as.data.frame(t(fitMeasures(fit))), file.path({r_quote(str(output_dir.resolve()))}, "CFA_fit_indices.csv"), row.names = FALSE)
write.csv(standardizedSolution(fit), file.path({r_quote(str(output_dir.resolve()))}, "CFA_standardized_solution.csv"), row.names = FALSE)
write.csv(reliability(fit), file.path({r_quote(str(output_dir.resolve()))}, "CFA_reliability_semTools.csv"))
'''
    (output_dir / "CFA_lavaan.R").write_text(cfa_script, encoding="utf-8")

    esem_script = f'''# Auto-generated by questionnaire_quality_pipeline.py
# ESEM is substantially more demanding than EFA/CFA. Run only when the report marks it feasible.
{mapping_comments}
library(readxl)
library(lavaan)

input_file <- {r_quote(str(Path(input_excel).resolve()))}
data <- read_excel(input_file, sheet = {r_quote(data_sheet)})
items <- c({ordered})
data[items] <- lapply(data[items], ordered)

model <- '
    {efa_left} =~ {all_items}
'

fit <- cfa(model, data = data, ordered = items, estimator = "WLSMV",
           rotation = "geomin", rotation.args = list(rstarts = 30))
summary(fit, fit.measures = TRUE, standardized = TRUE)
write.csv(as.data.frame(t(fitMeasures(fit))), file.path({r_quote(str(output_dir.resolve()))}, "ESEM_fit_indices.csv"), row.names = FALSE)
write.csv(standardizedSolution(fit), file.path({r_quote(str(output_dir.resolve()))}, "ESEM_standardized_solution.csv"), row.names = FALSE)
'''
    (output_dir / "ESEM_lavaan.R").write_text(esem_script, encoding="utf-8")


def maybe_run_r(output_dir: Path, cfg: Config, feasibility: dict[str, Any]) -> list[str]:
    messages = []
    if not cfg.run_r_if_available:
        return messages
    rscript = shutil.which("Rscript")
    if not rscript:
        return ["系統未找到Rscript，因此只產生lavaan程式碼，未自動執行。"]
    for filename, feasible_key in [("CFA_lavaan.R", "cfa_independent_feasible"), ("ESEM_lavaan.R", "esem_feasible")]:
        if not feasibility.get(feasible_key):
            continue
        try:
            proc = subprocess.run([rscript, str(output_dir / filename)], cwd=output_dir, capture_output=True, text=True, timeout=600)
            (output_dir / f"{filename}.log").write_text(proc.stdout + "\n" + proc.stderr, encoding="utf-8")
            messages.append(f"{filename}執行完成，return code={proc.returncode}。")
        except Exception as exc:
            messages.append(f"{filename}執行失敗：{type(exc).__name__}: {exc}")
    return messages


# -----------------------------------------------------------------------------
# Figures
# -----------------------------------------------------------------------------


def plot_scree(pa_p: dict[str, Any], pa_poly: dict[str, Any], output: Path) -> None:
    plt.figure(figsize=(8, 5))
    x = np.arange(1, len(pa_p["observed"]) + 1)
    plt.plot(x, pa_p["observed"], marker="o", label="Observed Pearson")
    plt.plot(x, pa_p["random_p95"], marker="o", linestyle="--", label="Pearson PA 95th")
    if np.all(np.isfinite(pa_poly.get("random_p95", np.array([np.nan])))):
        plt.plot(x, pa_poly["observed"], marker="s", label="Observed polychoric")
        plt.plot(x, pa_poly["random_p95"], marker="s", linestyle="--", label="Polychoric PA 95th")
    plt.axhline(1, linewidth=0.8)
    plt.xlabel("Factor number")
    plt.ylabel("Eigenvalue")
    plt.title("Scree plot and parallel analysis")
    plt.legend()
    plt.tight_layout()
    plt.savefig(output, dpi=180)
    plt.close()


def heatmap(df: pd.DataFrame, title: str, output: Path, vmin: float = -1, vmax: float = 1) -> None:
    width = max(7, 0.75 * df.shape[1] + 4)
    height = max(5, 0.35 * df.shape[0] + 2)
    plt.figure(figsize=(width, height))
    im = plt.imshow(df.values, aspect="auto", vmin=vmin, vmax=vmax)
    plt.xticks(range(df.shape[1]), df.columns, rotation=45, ha="right")
    plt.yticks(range(df.shape[0]), df.index)
    plt.colorbar(im, fraction=0.03, pad=0.03)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            plt.text(j, i, f"{df.iloc[i, j]:.2f}", ha="center", va="center", fontsize=8)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(output, dpi=180)
    plt.close()


def plot_reliability(rel: pd.DataFrame, output: Path) -> None:
    if rel.empty:
        return
    x = np.arange(len(rel))
    width = 0.2
    plt.figure(figsize=(max(7, len(rel) * 1.5), 4.5))
    plt.bar(x - 1.5 * width, rel["cronbach_alpha"], width, label="alpha")
    plt.bar(x - 0.5 * width, rel["ordinal_alpha"], width, label="ordinal alpha")
    plt.bar(x + 0.5 * width, rel["mcdonald_omega"], width, label="omega")
    plt.bar(x + 1.5 * width, rel["ordinal_omega"], width, label="ordinal omega")
    plt.axhline(0.70, linestyle="--", linewidth=1)
    plt.xticks(x, rel["factor"], rotation=30, ha="right")
    plt.ylim(0, 1.05)
    plt.ylabel("Reliability")
    plt.title("Internal consistency by factor")
    plt.legend()
    plt.tight_layout()
    plt.savefig(output, dpi=180)
    plt.close()


def plot_stability(stability: pd.DataFrame, output: Path) -> None:
    if stability.empty or "factor_assignment_stability" not in stability:
        return
    df = stability.sort_values("factor_assignment_stability")
    plt.figure(figsize=(8, max(4, len(df) * 0.3 + 1)))
    plt.barh(df["item"], df["factor_assignment_stability"])
    plt.axvline(0.80, linestyle="--", linewidth=1)
    plt.xlim(0, 1)
    plt.xlabel("Bootstrap factor-assignment stability")
    plt.title("Item stability")
    plt.tight_layout()
    plt.savefig(output, dpi=180)
    plt.close()


def best_model_for_k(
    comparison: pd.DataFrame,
    models: dict[tuple[int, str], dict[str, Any]],
    k: Optional[int],
) -> Optional[dict[str, Any]]:
    if not k:
        return None
    rows = comparison[(comparison["status"] == "ok") & (comparison["k"] == int(k))]
    if rows.empty:
        return None
    row = rows.sort_values(["decision_score", "penalty"]).iloc[0]
    return models.get((int(row["k"]), str(row["rotation"])))


def statistical_alternative_k(
    comparison: pd.DataFrame,
    evidence: pd.DataFrame,
    expected_k: Optional[int],
) -> Optional[int]:
    """Choose the strongest statistics-led alternative to the planned factor count."""
    pa = evidence[evidence["method"].str.contains("Parallel analysis", case=False, na=False)]
    supported = [int(v) for v in pa["suggested_factors"].dropna().tolist()]
    candidates = comparison[(comparison["status"] == "ok") & (comparison["k"].isin(supported))].copy()
    if expected_k:
        candidates = candidates[candidates["k"] != int(expected_k)]
    if candidates.empty:
        candidates = comparison[comparison["status"] == "ok"].copy()
        if expected_k:
            candidates = candidates[candidates["k"] != int(expected_k)]
    if candidates.empty:
        return None
    return int(candidates.sort_values(["decision_score", "penalty"]).iloc[0]["k"])


def section_factor_crosswalk(
    fit: Optional[dict[str, Any]],
    codebook: pd.DataFrame,
    items: list[str],
) -> pd.DataFrame:
    if fit is None or codebook.empty or "expected_factor" not in codebook.columns:
        return pd.DataFrame()
    cb = codebook.set_index("item")
    assigned = fit["item_table"].set_index("item")["assigned_factor"]
    rows = []
    for item in items:
        if item not in assigned.index or item not in cb.index:
            continue
        expected = cb.loc[item, "expected_factor"]
        if pd.isna(expected) or not str(expected).strip():
            continue
        rows.append({
            "item": item,
            "planned_section": str(expected).strip(),
            "empirical_factor": str(assigned.loc[item]),
        })
    if not rows:
        return pd.DataFrame()
    raw = pd.DataFrame(rows)
    counts = pd.crosstab(raw["planned_section"], raw["empirical_factor"])
    props = counts.div(counts.sum(axis=1), axis=0)
    out = counts.copy()
    out.columns = [f"count_{c}" for c in out.columns]
    for c in props.columns:
        out[f"proportion_{c}"] = props[c]
    out["dominant_empirical_factor"] = props.idxmax(axis=1)
    out["dominant_proportion"] = props.max(axis=1)
    out["planned_item_count"] = counts.sum(axis=1)
    return out.reset_index()


def merge_candidates_from_crosswalk(crosswalk: pd.DataFrame, minimum: float) -> pd.DataFrame:
    if crosswalk is None or crosswalk.empty:
        return pd.DataFrame()
    rows = []
    for factor, g in crosswalk.groupby("dominant_empirical_factor"):
        strong = g[g["dominant_proportion"] >= minimum]
        sections = strong["planned_section"].astype(str).tolist()
        if len(sections) >= 2:
            rows.append({
                "empirical_factor": factor,
                "planned_sections_combined": " + ".join(sections),
                "n_sections": len(sections),
                "minimum_assignment_proportion": strong["dominant_proportion"].min(),
                "interpretation": "統計上可能合併；概念上是否可合併仍需題目內容與理論審查。",
            })
    return pd.DataFrame(rows)


# -----------------------------------------------------------------------------
# Report text
# -----------------------------------------------------------------------------


def sample_size_interpretation(n: int, p: int) -> str:
    ratio = n / p
    if n < 50:
        return f"本研究僅有N={n}，每題約{ratio:.1f}人。此規模不足以把EFA視為穩定的量表定型證據，因此結果只能作探索性判斷；程式不切分樣本，也不將CFA/ESEM列為主要分析。"
    if n < 100:
        return f"本研究N={n}，仍屬偏小樣本。程式使用全部樣本進行EFA並加做bootstrap敏感度分析，不切分EFA/CFA樣本。"
    if n < 200:
        return f"本研究N={n}，可進行EFA，但再切分樣本通常會讓兩邊都偏小；CFA若進行，只能作同樣本診斷。"
    return f"本研究N={n}，樣本規模較有機會支持獨立CFA；是否切分仍取決於最終題數與模型參數。"


MANUSCRIPT_REFERENCES = [
    "American Educational Research Association, American Psychological Association, & National Council on Measurement in Education. (2014). Standards for Educational and Psychological Testing.",
    "Fabrigar, L. R., Wegener, D. T., MacCallum, R. C., & Strahan, E. J. (1999). Evaluating the use of exploratory factor analysis in psychological research. Psychological Methods, 4(3), 272–299.",
    "Garrido, L. E., Abad, F. J., & Ponsoda, V. (2013). A new look at Horn's parallel analysis with ordinal variables. Psychological Methods, 18(4), 454–474.",
    "Hayton, J. C., Allen, D. G., & Scarpello, V. (2004). Factor retention decisions in exploratory factor analysis: A tutorial on parallel analysis. Organizational Research Methods, 7(2), 191–205.",
    "MacCallum, R. C., Widaman, K. F., Zhang, S., & Hong, S. (1999). Sample size in factor analysis. Psychological Methods, 4(1), 84–99.",
    "McNeish, D. (2018). Thanks coefficient alpha, we'll take it from here. Psychological Methods, 23(3), 412–433.",
    "Mokkink, L. B., et al. (2010). The COSMIN checklist for assessing the methodological quality of studies on measurement properties. BMC Medical Research Methodology, 10, 22.",
    "Fornell, C., & Larcker, D. F. (1981). Evaluating structural equation models with unobservable variables and measurement error. Journal of Marketing Research, 18(1), 39–50.",
    "Henseler, J., Ringle, C. M., & Sarstedt, M. (2015). A new criterion for assessing discriminant validity in variance-based structural equation modeling. Journal of the Academy of Marketing Science, 43(1), 115–135.",
    "Saccenti, E., Hendriks, M. H. W. B., & Smilde, A. K. (2020). Corruption of the Pearson correlation coefficient by measurement error. Scientific Reports, 10, 438.",
    "Worthington, R. L., & Whittaker, T. A. (2006). Scale development research: A content analysis and recommendations for best practices. The Counseling Psychologist, 34(6), 806–838.",
]


def build_paper_summary(
    summary: dict[str, Any],
    corr_method: str,
    factor_evidence: pd.DataFrame,
    final_fit: dict[str, Any],
    deleted: pd.DataFrame,
    reliability: pd.DataFrame,
    validity: pd.DataFrame,
    htmt: pd.DataFrame,
    feasibility: dict[str, Any],
    expected_k: Optional[int] = None,
    alternative_k: Optional[int] = None,
    codebook: Optional[pd.DataFrame] = None,
) -> str:
    k = int(final_fit["metrics"]["k"])
    rotation = final_fit["metrics"]["rotation"]
    corr_name = "polychoric相關矩陣（polychoric correlation matrix）" if corr_method == "polychoric" else "Pearson相關矩陣（Pearson correlation matrix）"
    factor_ranges: list[str] = []
    factor_item_text: list[str] = []
    for _, row in reliability.iterrows():
        factor_ranges.append(
            f"{row['factor']}（{int(row['n_items'])}題，ω={fmt(row['mcdonald_omega'])}，ordinal ω={fmt(row['ordinal_omega'])}，α={fmt(row['cronbach_alpha'])}）"
        )
        factor_item_text.append(f"{row['factor']}：{row['items']}")
    deletion_text = "未自動刪除題目" if deleted.empty else f"依序刪除{len(deleted)}題（{', '.join(deleted['item'].astype(str))}）"
    loading_min = final_fit["item_table"]["primary_loading"].min()
    loading_max = final_fit["item_table"]["primary_loading"].max()
    ave_text = ""
    if not validity.empty and validity["average_variance_extracted_AVE"].notna().any():
        ave_text = f"探索性AVE介於{validity['average_variance_extracted_AVE'].min():.3f}至{validity['average_variance_extracted_AVE'].max():.3f}。"
    htmt_text = ""
    if htmt is not None and not htmt.empty and htmt.shape[0] > 1:
        values = offdiag_values(htmt.values.astype(float))
        if np.isfinite(values).any():
            htmt_text = f"探索性HTMT最大值為{np.nanmax(values):.3f}。"

    common = (
        "【共同方法段落】\n"
        f"本研究以{corr_name}進行探索性因素分析（exploratory factor analysis, EFA），採主軸因素法（principal axis factoring, PAF），"
        "並綜合平行分析（parallel analysis）、Velicer最小平均偏相關法（minimum average partial, MAP）、陡坡圖（scree plot）、預定理論構面及因素解可解釋性判斷因素數。"
        f"主要方案保留{k}個因素，採{rotation}轉軸。刪題時同時考量主要因素負荷（primary factor loading）、交叉負荷（cross-loading）、共同性（communality）、個別取樣適切性（measure of sampling adequacy, MSA）、理論歸屬及每因素最低題數；{deletion_text}。"
        f"正式計分題目共{len(final_fit['item_table'])}題，配置為：{'；'.join(factor_item_text)}。主要負荷介於{loading_min:.3f}至{loading_max:.3f}，RMSR={final_fit['metrics']['rmsr']:.3f}。"
        f"各分量表信度為：{'；'.join(factor_ranges)}。{ave_text}{htmt_text}"
        "這些結果提供信度證據與基於內部結構的初步效度證據（preliminary validity evidence based on internal structure），不等同於已建立完整構念效度、效標關聯效度、跨群組測量恆等性或重測穩定性。"
    )

    scenarios = [common, "", "【情境A：採用統計主要方案】",
        f"根據上述結果，本研究以{k}因素方案計算分量表分數並進行後續分析。後續相關或迴歸係數應連同各分量表信度與本樣本的探索性結構限制一併解讀。{feasibility['decision']}"]

    if expected_k and alternative_k and alternative_k != expected_k:
        scenarios += ["", f"【情境B：基於理論仍保留原定{expected_k}個section】",
            f"雖然平行分析及候選模型比較提出{alternative_k}因素作為統計替代方案，本研究仍依問卷發展階段的理論定義保留原定{expected_k}個section作理論導向計分（theory-driven scoring）。"
            f"因此，不將本次EFA描述為已證實{expected_k}因素結構，而是將{alternative_k}因素解列為敏感度分析（sensitivity analysis），並將原定分量表間的區分視為尚待獨立樣本CFA或ESEM驗證。若以原定分量表進行相關或迴歸分析，相關結果應標示為探索性，並避免對高度重疊分量表的獨立效果作強結論。"]
        scenarios += ["", "【情境C：概念上不能合併section】",
            "若統計替代解將兩個section合併，但理論上兩構念不可合併，本研究保留兩者的原定計分，並將結果解讀為目前題目措辭或樣本未能充分區辨兩構念，而非兩構念在概念上相同。應同時報告替代解、因素相關、HTMT及題目交叉負荷，並在限制中說明區辨效度證據不足。"]

    scenarios += ["", "【情境D：保留內容必要但統計較弱的題目】",
        "若某題在因素負荷、共同性或題目總分相關上較弱，但它涵蓋其餘題目未測量的必要內容，可將該題標為內容保護題（content-essential/protected item）並暫時保留。論文應同時呈現含題與不含題的敏感度結果，說明保留的內容效度理由，並將該題列為未來改寫、認知訪談及新樣本再驗證的優先項目。"]
    scenarios += ["", "【後續相關與迴歸分析的使用聲明】",
        "量表分數能否用於相關（correlation）或迴歸（regression）不是由單一門檻自動決定。若因素結構可解釋、分量表信度可接受且計分方式與理論一致，可進行後續分析，但仍需報告測量限制。若信度低，測量誤差通常會降低統計功效並使部分關聯衰減；若構面區辨不足，將高度重疊的分量表同時放入迴歸可能造成係數不穩定或多重共線性。若結構與信度均明顯不佳，不建議把分量表分數當作已建立的構念進行確認性推論；最多可作探索性／敏感度分析。"]
    scenarios += ["", "【參考文獻】"] + [f"- {r}" for r in MANUSCRIPT_REFERENCES]
    return "\n".join(scenarios)


def build_decision_summary(analysis: dict[str, Any]) -> str:
    n = int(analysis["data_summary"]["n_analysis"])
    expected_k = analysis.get("expected_k")
    final_k = int(analysis["final_fit"]["metrics"]["k"])
    alternative_k = analysis.get("alternative_k")
    deleted = analysis.get("deleted_items", pd.DataFrame())
    rel = analysis.get("reliability", pd.DataFrame())
    omega_min = float(rel["mcdonald_omega"].min()) if not rel.empty and rel["mcdonald_omega"].notna().any() else np.nan
    critical_structure = (
        analysis["final_fit"]["metrics"].get("heywood_count", 0) > 0
        or analysis["final_fit"]["metrics"].get("factor_below_three_items", 0) > 0
        or analysis["final_fit"]["metrics"].get("maximum_factor_correlation", 0) >= analysis["config"].planned_factor_correlation_critical
    )
    if critical_structure:
        use = "目前不宜把所有分量表視為已建立且可清楚區分的構念；應先修訂，後續推論最多作探索性分析。"
    elif n < analysis["config"].small_sample_threshold:
        use = "可依理論作暫時或探索性使用，但樣本偏小，不能把本次結果視為量表已定型。"
    elif np.isfinite(omega_min) and omega_min < 0.60:
        use = "可作有限度的探索性使用，但至少一個分量表信度偏低，後續關聯可能受測量誤差影響。"
    else:
        use = "可用於本樣本的後續分析，但仍應將目前證據描述為信度與基於內部結構的初步效度證據。"
    structure = f"原定{expected_k}個section；主要工作方案為{final_k}因素" if expected_k else f"主要工作方案為{final_k}因素"
    if alternative_k and alternative_k != final_k:
        structure += f"，另保留{alternative_k}因素統計替代方案"
    deletion = "未自動刪題" if deleted is None or deleted.empty else f"自動刪除{len(deleted)}題（{', '.join(deleted['item'].astype(str))}）"
    reliability = f"各分量表最低McDonald’s ω為{fmt(omega_min)}" if np.isfinite(omega_min) else "部分分量表無法穩定估計omega"
    return f"使用判斷：{use} 結構判斷：{structure}。題目處理：{deletion}。信度摘要：{reliability}。"


# -----------------------------------------------------------------------------
# Excel and report generation
# -----------------------------------------------------------------------------


def flatten_dict(d: dict[str, Any], prefix: str = "") -> list[dict[str, Any]]:
    rows = []
    for k, v in d.items():
        key = f"{prefix}.{k}" if prefix else k
        if isinstance(v, dict):
            rows.extend(flatten_dict(v, key))
        else:
            rows.append({"metric": key, "value": json.dumps(v, ensure_ascii=False) if isinstance(v, (list, tuple)) else v})
    return rows


def write_side_excel(path: Path, sheets: dict[str, pd.DataFrame], notes: list[str]) -> None:
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame({"說明": notes}).to_excel(writer, sheet_name="README", index=False)
        for name, df in sheets.items():
            if df is None:
                continue
            safe = sanitize_sheet_name(name)
            if isinstance(df, pd.Series):
                df = df.to_frame()
            df.to_excel(writer, sheet_name=safe, index=True if df.index.name or not isinstance(df.index, pd.RangeIndex) else False)
        wb = writer.book
        for ws in wb.worksheets:
            ws.freeze_panes = "A2"
            for cell in ws[1]:
                font = copy(cell.font)
                font.bold = True
                cell.font = font
            for col_cells in ws.columns:
                max_len = 0
                col_letter = col_cells[0].column_letter
                for cell in col_cells[:200]:
                    value = "" if cell.value is None else str(cell.value)
                    max_len = max(max_len, len(value))
                ws.column_dimensions[col_letter].width = min(max(max_len + 2, 10), 45)


def add_docx_table(doc: Document, df: pd.DataFrame, max_rows: int = 40, digits: int = 3) -> None:
    if df is None or df.empty:
        doc.add_paragraph("無可呈現結果。")
        return
    shown = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.style = "Table Grid"
    for j, col in enumerate(shown.columns):
        table.rows[0].cells[j].text = str(col)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = fmt(value, digits) if isinstance(value, (float, np.floating)) else str(value)
    if len(df) > max_rows:
        doc.add_paragraph(f"主報告僅顯示前{max_rows}列；完整結果見Questionnaire_quality_details.xlsx。")


def make_docx_report(
    path: Path,
    analysis: dict[str, Any],
    figure_paths: dict[str, Path],
) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"].font.size = Pt(10.5)
    title = doc.add_heading("問卷品質、測量結構與信效度評估報告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"資料檔：{analysis['input_file']}")

    doc.add_heading("一、結論摘要", level=1)
    doc.add_paragraph(build_decision_summary(analysis))
    doc.add_paragraph("此處先回答問卷能否使用、因素數、刪題與信度；完整方法、替代決策與可改寫至論文的多情境段落見 paper_ready_summary.txt。")

    doc.add_heading("二、資料與樣本適切性", level=1)
    doc.add_paragraph(analysis["sample_interpretation"])
    doc.add_paragraph(
        f"原始樣本數為{analysis['data_summary']['n_original']}，排除遺漏超過設定門檻的作答後，分析樣本為{analysis['data_summary']['n_analysis']}。"
        f"題目數為{analysis['data_summary']['n_items']}，整體遺漏比例為{analysis['data_summary']['overall_missing_fraction']:.1%}。"
    )
    flagged = analysis["data_check_main"].copy()
    flagged = flagged[(flagged["missing_fraction"] > 0.10) | (flagged["zero_variance"]) | (flagged["skewness"].abs() > 2)]
    if flagged.empty:
        doc.add_paragraph("逐題資料檢查未發現需要在主報告中特別列示的嚴重問題；完整描述統計見Excel附檔。")
    else:
        doc.add_paragraph("下表僅列出需要注意的題目；其餘逐題數據見Excel附檔。")
        add_docx_table(doc, flagged[["item", "missing_fraction", "sd", "unique_categories", "skewness", "floor_fraction", "ceiling_fraction"]], max_rows=15)

    doc.add_heading("三、相關矩陣的選擇", level=1)
    for reason in analysis["correlation_reasons"]:
        doc.add_paragraph(reason, style="List Bullet")
    doc.add_paragraph(f"主分析採用：{'polychoric相關' if analysis['primary_correlation_method']=='polychoric' else 'Pearson相關'}。另一種相關矩陣保留在附檔作敏感度分析。")
    doc.add_paragraph(
        f"主分析KMO={analysis['factorability']['kmo_total']:.3f}；Bartlett球形檢定χ²({analysis['factorability']['bartlett_df']})={analysis['factorability']['bartlett_chi2']:.3f}，p={analysis['factorability']['bartlett_p']:.4g}。"
    )

    doc.add_heading("四、因素數判斷", level=1)
    doc.add_paragraph("因素數不是只依特徵值大於1決定，而是綜合平行分析、MAP、理論預定構面與候選模型的簡單結構。")
    add_docx_table(doc, analysis["factor_count_evidence"], max_rows=10)
    doc.add_paragraph("表中統計方法若不一致，不以多數決自動決定；原定構面會被優先保留，除非出現因素題數不足、Heywood case、因素高度重疊、殘差偏高或大量低負荷／交叉負荷等強烈反證。")
    comparison_rows = []
    if analysis.get("planned_fit") is not None:
        m = analysis["planned_fit"]["metrics"]
        comparison_rows.append({"方案": f"原定{analysis.get('expected_k')}因素", "RMSR": m.get("rmsr"), "低負荷題": m.get("low_loading_count"), "交叉負荷題": m.get("cross_loading_count"), "因素題數不足": m.get("factor_below_three_items"), "最大因素相關": m.get("maximum_factor_correlation")})
    if analysis.get("statistical_alternative_fit") is not None:
        m = analysis["statistical_alternative_fit"]["metrics"]
        comparison_rows.append({"方案": f"統計替代{analysis.get('alternative_k')}因素", "RMSR": m.get("rmsr"), "低負荷題": m.get("low_loading_count"), "交叉負荷題": m.get("cross_loading_count"), "因素題數不足": m.get("factor_below_three_items"), "最大因素相關": m.get("maximum_factor_correlation")})
    if comparison_rows:
        doc.add_paragraph("下表僅比較會影響決策的核心指標；完整負荷矩陣與section對照見HTML及Excel。")
        add_docx_table(doc, pd.DataFrame(comparison_rows), max_rows=5)

    doc.add_heading("五、萃取法與轉軸判斷", level=1)
    doc.add_paragraph("主分析使用主軸因素法，因問卷資料為Likert題且樣本通常不大；只有在樣本與分布條件較好時，最大概似法才列為附加敏感度分析。")
    for reason in analysis["model_reasons"]:
        doc.add_paragraph(reason, style="List Bullet")
    add_docx_table(doc, analysis["candidate_models"][[c for c in ["k", "rotation", "rmsr", "mean_primary_loading", "cross_loading_count", "low_loading_count", "factor_below_three_items", "maximum_factor_correlation", "decision_score", "status"] if c in analysis["candidate_models"].columns]], max_rows=20)

    doc.add_heading("六、刪題迭代", level=1)
    if analysis["deleted_items"].empty:
        doc.add_paragraph("沒有題目同時達到嚴重統計問題且符合可刪除條件，因此未自動刪題。")
    else:
        doc.add_paragraph("每次僅刪除一題，刪除後重新估計因素模型。刪題理由如下：")
        add_docx_table(doc, analysis["deleted_items"][["iteration", "item", "reason", "primary_loading", "secondary_loading", "loading_gap", "communality", "msa", "expected_factor"]])

    doc.add_heading("七、最終因素結構", level=1)
    doc.add_paragraph(
        f"最終保留{len(analysis['final_items'])}題、{analysis['final_fit']['metrics']['k']}個因素；轉軸為{analysis['final_fit']['metrics']['rotation']}，RMSR={analysis['final_fit']['metrics']['rmsr']:.3f}。"
    )
    final_summary = analysis["final_fit"]["item_table"][["item", "assigned_factor", "primary_loading", "secondary_loading", "loading_gap", "communality"]].copy()
    doc.add_paragraph("下表為最終題目歸屬與核心診斷；完整pattern matrix與structure matrix見HTML與Excel。")
    add_docx_table(doc, final_summary, max_rows=25)

    doc.add_heading("八、信度與效度證據", level=1)
    doc.add_paragraph("信度分別依各因素計算，不以多構面問卷總分的α作為單一品質指標。主要報告McDonald’s omega，並附Cronbach’s alpha與序位版本。")
    add_docx_table(doc, analysis["reliability"])
    add_docx_table(doc, analysis["validity"])
    if not analysis["htmt"].empty:
        values = offdiag_values(analysis["htmt"].values.astype(float))
        max_htmt = float(np.nanmax(values)) if np.isfinite(values).any() else np.nan
        doc.add_paragraph(f"HTMT最大值為{fmt(max_htmt)}；完整矩陣見HTML與Excel。")
    doc.add_paragraph("本程式能提供內部結構、聚合與區辨性相關證據；若資料中沒有外部效標、已知群組或重測資料，不能自動產生效標關聯效度、已知群組效度或重測信度。")

    doc.add_heading("九、小樣本穩定性", level=1)
    doc.add_paragraph("Bootstrap用來檢查題目在重抽樣下是否仍落在相同因素。Word僅保留文字結論；逐題穩定性與圖形見HTML與Excel。")
    if analysis["stability"] is not None and not analysis["stability"].empty:
        stable_col = "same_factor_fraction" if "same_factor_fraction" in analysis["stability"].columns else None
        if stable_col:
            doc.add_paragraph(f"題目因素歸屬穩定比例中位數為{analysis['stability'][stable_col].median():.1%}。")

    doc.add_heading("十、CFA／ESEM可行性", level=1)
    doc.add_paragraph(analysis["feasibility"]["decision"])
    add_docx_table(doc, pd.DataFrame(flatten_dict(analysis["feasibility"])))
    doc.add_paragraph("程式已產生CFA_lavaan.R與ESEM_lavaan.R。Likert題在lavaan中以ordered變項及WLSMV估計；只有報告判定樣本足夠時才建議執行。")

    doc.add_heading("十一、論文報告與參考文獻", level=1)
    doc.add_paragraph("多情境論文段落、後續相關／迴歸的使用聲明及完整參考文獻已保存於 paper_ready_summary.txt；方法背景與術語說明見 README.md 與 REFERENCES.md。")

    doc.add_heading("十二、限制與使用原則", level=1)
    limitations = [
        "EFA是樣本依賴的；尤其N約20至30時，結果只能作探索性證據，不宜宣稱量表已完成穩定驗證。",
        "自動刪題不能取代內容效度判斷。刪題前仍應閱讀題目文字，確認是否刪除了構念的重要內容。",
        "同一份資料先刪題再做CFA，不是獨立驗證。正式研究宜另蒐集樣本。",
        "alpha或omega高不代表單向度，也可能來自題目重複；因此應同時判讀因素負荷、共同性、HTMT及題目內容。",
        "低信度與多重共線性是不同問題；但構面區辨不足時，將高度重疊分量表同時放入迴歸可能造成多重共線性。",
    ]
    for x in limitations:
        doc.add_paragraph(x, style="List Bullet")
    doc.save(path)


HTML_TEMPLATE = Template(r"""
<!doctype html><html lang="zh-Hant"><head><meta charset="utf-8">
<title>問卷品質、測量結構與信效度評估報告</title>
<style>
body{font-family:-apple-system,BlinkMacSystemFont,"Microsoft JhengHei",sans-serif;max-width:1100px;margin:auto;padding:30px;line-height:1.65;color:#222}
h1,h2{color:#17365d} table{border-collapse:collapse;width:100%;font-size:13px;margin:12px 0 24px} th,td{border:1px solid #bbb;padding:6px;text-align:left} th{background:#eaf0f8;position:sticky;top:0}.note{background:#fff7d6;border-left:5px solid #d6a700;padding:12px}.good{background:#eef8ee;border-left:5px solid #4b8f4b;padding:12px}img{max-width:100%;height:auto}.small{font-size:12px;color:#555}</style></head><body>
<h1>問卷品質、測量結構與信效度評估報告</h1>
<div class="good"><b>可直接改寫至論文：</b><br>{{ paper_summary }}</div>
<h2>1. 樣本與資料品質</h2><p>{{ sample_interpretation }}</p>{{ data_table|safe }}
<h2>2. 相關矩陣與因素分析適切性</h2><ul>{% for x in correlation_reasons %}<li>{{x}}</li>{% endfor %}</ul>
<p>KMO={{kmo}}；Bartlett χ²({{bart_df}})={{bart_chi}}, p={{bart_p}}</p>
<h2>3. 因素數與候選模型</h2><p>因素數證據若不一致，原定構面與統計替代方案會同時保留。下列圖表用於比較，而非以單一門檻自動裁決。</p>{{ factor_count_table|safe }}{% if scree %}<img src="{{scree}}">{% endif %}{{ candidate_table|safe }}
<h2>4. 原定section與統計替代方案</h2>
<p>原定方案回答「依理論固定section數時資料表現如何」；統計替代方案回答「若讓資料提出較精簡或不同結構，題目如何重新聚集」。</p>
{% if planned_table %}<h3>原定方案</h3>{{ planned_table|safe }}{{ planned_crosswalk|safe }}{% if planned_fig %}<img src="{{planned_fig}}">{% endif %}{% endif %}
{% if alternative_table %}<h3>統計替代方案</h3>{{ alternative_table|safe }}{{ alternative_crosswalk|safe }}{{ merge_candidates|safe }}{% if alternative_fig %}<img src="{{alternative_fig}}">{% endif %}{% endif %}
<h2>5. 刪題紀錄</h2>{{ deleted_table|safe }}
<h2>6. 最終工作方案</h2>{{ loading_table|safe }}{% if loadings %}<img src="{{loadings}}">{% endif %}{% if phi %}<img src="{{phi}}">{% endif %}
<h2>7. 信度與初步效度證據</h2>{{ reliability_table|safe }}{{ validity_table|safe }}{{ htmt_table|safe }}{% if reliability_fig %}<img src="{{reliability_fig}}">{% endif %}
<h2>8. 穩定性</h2>{{ stability_table|safe }}{% if stability_fig %}<img src="{{stability_fig}}">{% endif %}
<h2>9. CFA／ESEM</h2><p>{{ feasibility }}</p>
<div class="note">完整的替代相關矩陣、候選模型、刪題分支、pattern/structure matrix與殘差矩陣，請見 Questionnaire_quality_details.xlsx。</div>
</body></html>
""")


def image_data_uri(path: Optional[Path]) -> str:
    if not path or not path.exists():
        return ""
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{data}"


def make_html_report(path: Path, analysis: dict[str, Any], figures: dict[str, Path]) -> None:
    html = HTML_TEMPLATE.render(
        paper_summary=analysis["paper_summary"],
        sample_interpretation=analysis["sample_interpretation"],
        data_table=analysis["data_check_main"].head(40).to_html(index=False, float_format=lambda x: f"{x:.3f}"),
        correlation_reasons=analysis["correlation_reasons"],
        kmo=fmt(analysis["factorability"]["kmo_total"]),
        bart_df=analysis["factorability"]["bartlett_df"],
        bart_chi=fmt(analysis["factorability"]["bartlett_chi2"]),
        bart_p=f"{analysis['factorability']['bartlett_p']:.4g}",
        factor_count_table=analysis["factor_count_evidence"].to_html(index=False),
        candidate_table=analysis["candidate_models"].head(20).to_html(index=False, float_format=lambda x: f"{x:.3f}"),
        planned_table=(analysis["planned_fit"]["item_table"].to_html(index=False, float_format=lambda x: f"{x:.3f}") if analysis.get("planned_fit") is not None else ""),
        planned_crosswalk=(analysis["planned_crosswalk"].to_html(index=False, float_format=lambda x: f"{x:.3f}") if analysis.get("planned_crosswalk") is not None and not analysis["planned_crosswalk"].empty else ""),
        alternative_table=(analysis["statistical_alternative_fit"]["item_table"].to_html(index=False, float_format=lambda x: f"{x:.3f}") if analysis.get("statistical_alternative_fit") is not None else ""),
        alternative_crosswalk=(analysis["alternative_crosswalk"].to_html(index=False, float_format=lambda x: f"{x:.3f}") if analysis.get("alternative_crosswalk") is not None and not analysis["alternative_crosswalk"].empty else ""),
        merge_candidates=(analysis["merge_candidates"].to_html(index=False, float_format=lambda x: f"{x:.3f}") if analysis.get("merge_candidates") is not None and not analysis["merge_candidates"].empty else "<p>未形成明確section合併候選，或合併比例未達設定門檻。</p>"),
        deleted_table=(analysis["deleted_items"].to_html(index=False, float_format=lambda x: f"{x:.3f}") if not analysis["deleted_items"].empty else "<p>未自動刪題。</p>"),
        loading_table=analysis["final_fit"]["loadings"].reset_index().rename(columns={"index": "item"}).to_html(index=False, float_format=lambda x: f"{x:.3f}"),
        reliability_table=analysis["reliability"].to_html(index=False, float_format=lambda x: f"{x:.3f}"),
        validity_table=analysis["validity"].to_html(index=False, float_format=lambda x: f"{x:.3f}"),
        htmt_table=(analysis["htmt"].to_html(float_format=lambda x: f"{x:.3f}") if not analysis["htmt"].empty else ""),
        stability_table=analysis["stability"].to_html(index=False, float_format=lambda x: f"{x:.3f}"),
        feasibility=analysis["feasibility"]["decision"],
        scree=image_data_uri(figures.get("scree")),
        loadings=image_data_uri(figures.get("loadings")),
        phi=image_data_uri(figures.get("phi")),
        reliability_fig=image_data_uri(figures.get("reliability")),
        stability_fig=image_data_uri(figures.get("stability")),
        planned_fig=image_data_uri(figures.get("planned_loadings")),
        alternative_fig=image_data_uri(figures.get("alternative_loadings")),
    )
    path.write_text(html, encoding="utf-8")


# -----------------------------------------------------------------------------
# Main analysis
# -----------------------------------------------------------------------------


def run_analysis(input_file: str, output_dir: str, config_file: Optional[str] = None) -> dict[str, Any]:
    cfg = Config.from_yaml(config_file)
    out = ensure_dir(Path(output_dir))
    figdir = ensure_dir(out / "figures")

    raw, codebook, items, input_messages = read_input(input_file, cfg)
    data_prepared, codebook, reverse_messages = prepare_items(raw, codebook, items, cfg)
    clean, item_diag, frequencies, row_diag, data_summary = data_checks(data_prepared, items, cfg)

    zero_items = data_summary["zero_variance_items"]
    if zero_items:
        items = [x for x in items if x not in zero_items]
        clean = clean[items]
    if len(items) < 4:
        raise RuntimeError("移除零變異題目後，題目不足以進行EFA。")

    pearson_corr, pearson_diag = pearson_correlation(clean[items])
    poly_corr, poly_pairs, poly_diag = polychoric_correlation(clean[items])
    primary_method, corr_reasons = choose_correlation_method(item_diag[item_diag["item"].isin(items)], len(clean), poly_diag, cfg)
    primary_corr = poly_corr if primary_method == "polychoric" else pearson_corr
    alternative_corr = pearson_corr if primary_method == "polychoric" else poly_corr

    msa_values, kmo_total = kmo_from_corr(primary_corr.values)
    bart_chi, bart_df, bart_p = bartlett_from_corr(primary_corr.values, len(clean))
    factorability = {
        "kmo_total": kmo_total,
        "bartlett_chi2": bart_chi,
        "bartlett_df": bart_df,
        "bartlett_p": bart_p,
        "item_msa": dict(zip(items, msa_values)),
    }

    expected_levels, _ = planned_factor_info(codebook, items)
    if cfg.expected_factors is None and expected_levels:
        cfg.expected_factors = len(expected_levels)

    evidence, pa_p, pa_poly, map_p, map_poly = factor_count_evidence(clean[items], pearson_corr, poly_corr, cfg)
    primary_comparison, primary_models = compare_models(primary_corr, len(clean), items, evidence, cfg, codebook)
    chosen_initial, model_reasons = choose_model(primary_comparison, primary_models, cfg, codebook, items, len(clean), evidence)
    planned_levels, _ = planned_factor_info(codebook, items)
    expected_k = cfg.expected_factors or (len(planned_levels) if planned_levels else None)
    planned_fit = best_model_for_k(primary_comparison, primary_models, expected_k)
    alternative_k = statistical_alternative_k(primary_comparison, evidence, expected_k)
    statistical_alternative_fit = best_model_for_k(primary_comparison, primary_models, alternative_k)
    planned_crosswalk = section_factor_crosswalk(planned_fit, codebook, items)
    alternative_crosswalk = section_factor_crosswalk(statistical_alternative_fit, codebook, items)
    merge_candidates = merge_candidates_from_crosswalk(alternative_crosswalk, cfg.merge_assignment_minimum)

    # Complete alternative-correlation model comparison for side results.
    alt_comparison, alt_models = compare_models(alternative_corr, len(clean), items, evidence, cfg, codebook)

    refinement = iterative_item_refinement(
        clean,
        items,
        codebook,
        cfg,
        primary_method,
        chosen_initial["metrics"]["k"],
        chosen_initial["metrics"]["rotation"],
    )
    final_items = refinement["final_items"]
    final_fit = refinement["final_fit"]
    if final_fit is None:
        final_fit = chosen_initial
        final_items = items
    # Recompute both final correlation matrices and final fit using the final chosen settings.
    final_pearson, final_pearson_diag = pearson_correlation(clean[final_items])
    final_poly, final_poly_pairs, final_poly_diag = polychoric_correlation(clean[final_items])
    final_primary_corr = final_poly if primary_method == "polychoric" else final_pearson
    levels_final, mapping_final = planned_factor_info(codebook, final_items)
    final_fit = fit_factor_model(
        final_primary_corr,
        len(clean),
        final_fit["metrics"]["k"],
        "pa",
        final_fit["metrics"]["rotation"],
        final_items,
        levels_final,
        mapping_final,
    )

    # Final sensitivity model using the alternative correlation matrix.
    try:
        final_alt_corr = final_pearson if primary_method == "polychoric" else final_poly
        final_alt_fit = fit_factor_model(
            final_alt_corr,
            len(clean),
            final_fit["metrics"]["k"],
            "pa",
            final_fit["metrics"]["rotation"],
            final_items,
            levels_final,
            mapping_final,
        )
    except Exception:
        final_alt_fit = None

    reliability, validity, htmt = reliability_validity(clean, final_fit, final_items, final_poly, cfg)
    planned_reliability = planned_validity = planned_htmt = pd.DataFrame()
    alternative_reliability = alternative_validity = alternative_htmt = pd.DataFrame()
    if planned_fit is not None:
        planned_reliability, planned_validity, planned_htmt = reliability_validity(clean, planned_fit, items, poly_corr, cfg)
    if statistical_alternative_fit is not None:
        alternative_reliability, alternative_validity, alternative_htmt = reliability_validity(clean, statistical_alternative_fit, items, poly_corr, cfg)
    stability = bootstrap_stability(clean, final_fit, final_items, primary_method, cfg)

    factor_counts = Counter(final_fit["item_table"]["assigned_factor"])
    min_factor_items = min(factor_counts.values()) if factor_counts else 0
    feasibility = confirmatory_feasibility(len(clean), len(final_items), final_fit["metrics"]["k"], min_factor_items)
    generate_lavaan_scripts(out, input_file, cfg.data_sheet, final_fit, final_items, feasibility)
    r_messages = maybe_run_r(out, cfg, feasibility)

    paper_summary = build_paper_summary(data_summary, primary_method, evidence, final_fit, refinement["deleted_items"], reliability, validity, htmt, feasibility, expected_k=expected_k, alternative_k=alternative_k, codebook=codebook)
    sample_interpretation = sample_size_interpretation(len(clean), len(items))

    # Figures
    figures = {
        "scree": figdir / "01_scree_parallel.png",
        "loadings": figdir / "02_final_loadings.png",
        "phi": figdir / "03_factor_correlations.png",
        "reliability": figdir / "04_reliability.png",
        "stability": figdir / "05_bootstrap_stability.png",
        "planned_loadings": figdir / "06_planned_solution_loadings.png",
        "alternative_loadings": figdir / "07_statistical_alternative_loadings.png",
    }
    plot_scree(pa_p, pa_poly, figures["scree"])
    heatmap(final_fit["loadings"], "Final pattern matrix", figures["loadings"], -1, 1)
    heatmap(final_fit["phi"], "Factor correlations", figures["phi"], -1, 1)
    plot_reliability(reliability, figures["reliability"])
    plot_stability(stability, figures["stability"])
    if planned_fit is not None:
        heatmap(planned_fit["loadings"], f"Planned {expected_k}-factor solution", figures["planned_loadings"], -1, 1)
    if statistical_alternative_fit is not None:
        heatmap(statistical_alternative_fit["loadings"], f"Statistical alternative {alternative_k}-factor solution", figures["alternative_loadings"], -1, 1)

    analysis = {
        "input_file": str(Path(input_file).resolve()),
        "config": cfg,
        "input_messages": input_messages + reverse_messages + r_messages,
        "data_summary": data_summary,
        "data_check_main": item_diag,
        "frequencies": frequencies,
        "row_diagnostics": row_diag,
        "primary_correlation_method": primary_method,
        "correlation_reasons": corr_reasons,
        "pearson_corr": pearson_corr,
        "pearson_diag": pearson_diag,
        "poly_corr": poly_corr,
        "poly_pairs": poly_pairs,
        "poly_diag": poly_diag,
        "factorability": factorability,
        "factor_count_evidence": evidence,
        "pa_pearson": pa_p,
        "pa_poly": pa_poly,
        "map_pearson": map_p,
        "map_poly": map_poly,
        "candidate_models": primary_comparison,
        "alternative_candidate_models": alt_comparison,
        "expected_k": expected_k,
        "alternative_k": alternative_k,
        "planned_fit": planned_fit,
        "statistical_alternative_fit": statistical_alternative_fit,
        "planned_crosswalk": planned_crosswalk,
        "alternative_crosswalk": alternative_crosswalk,
        "merge_candidates": merge_candidates,
        "planned_reliability": planned_reliability,
        "planned_validity": planned_validity,
        "planned_htmt": planned_htmt,
        "alternative_reliability": alternative_reliability,
        "alternative_validity": alternative_validity,
        "alternative_htmt": alternative_htmt,
        "model_reasons": model_reasons,
        "initial_fit": chosen_initial,
        "iteration_log": refinement["iteration_log"],
        "deleted_items": refinement["deleted_items"],
        "alternative_deletions": refinement["alternative_deletions"],
        "final_items": final_items,
        "final_fit": final_fit,
        "final_alt_fit": final_alt_fit,
        "final_pearson": final_pearson,
        "final_poly": final_poly,
        "reliability": reliability,
        "validity": validity,
        "htmt": htmt,
        "stability": stability,
        "feasibility": feasibility,
        "paper_summary": paper_summary,
        "sample_interpretation": sample_interpretation,
    }

    # Side workbook
    msa_table = pd.DataFrame({"item": items, "MSA": msa_values})
    factor_count_detail = pd.DataFrame({
        "factor_number": np.arange(1, len(pa_p["observed"]) + 1),
        "observed_eigenvalue_pearson": pa_p["observed"],
        "random_95th_pearson": pa_p["random_p95"],
        "observed_eigenvalue_polychoric": pa_poly["observed"],
        "random_95th_polychoric": pa_poly["random_p95"],
    })
    final_item_diagnostics = final_fit["item_table"].copy()
    final_item_diagnostics["MSA"] = pd.Series(kmo_from_corr(final_primary_corr.values)[0], index=final_items).values
    sensitivity = pd.DataFrame()
    if final_alt_fit is not None:
        sensitivity = pd.DataFrame({
            "item": final_items,
            "primary_method_factor": final_fit["item_table"]["assigned_factor"],
            "alternative_method_factor": final_alt_fit["item_table"]["assigned_factor"],
            "primary_method_loading": final_fit["item_table"]["primary_loading"],
            "alternative_method_loading": final_alt_fit["item_table"]["primary_loading"],
        })
        sensitivity["same_factor_assignment"] = sensitivity["primary_method_factor"] == sensitivity["alternative_method_factor"]

    side_sheets = {
        "Config": pd.DataFrame(flatten_dict(asdict(cfg))),
        "Data_Check": item_diag,
        "Response_Frequencies": frequencies,
        "Row_Diagnostics": row_diag,
        "Pearson_Correlation": pearson_corr,
        "Polychoric_Correlation": poly_corr,
        "Polychoric_Pairs": poly_pairs,
        "Correlation_Diagnostics": pd.DataFrame(flatten_dict({"pearson": pearson_diag, "polychoric": poly_diag})),
        "KMO_MSA": msa_table,
        "Factor_Count_Evidence": evidence,
        "Parallel_Analysis": factor_count_detail,
        "MAP_Pearson": map_p,
        "MAP_Polychoric": map_poly,
        "Candidate_Models_Primary": primary_comparison,
        "Candidate_Models_Alternative": alt_comparison,
        "Planned_Pattern_Matrix": (planned_fit["loadings"] if planned_fit is not None else pd.DataFrame()),
        "Planned_Factor_Correlations": (planned_fit["phi"] if planned_fit is not None else pd.DataFrame()),
        "Planned_Section_Crosswalk": planned_crosswalk,
        "Planned_Reliability": planned_reliability,
        "Planned_Validity": planned_validity,
        "Planned_HTMT": planned_htmt,
        "Alternative_Pattern_Matrix": (statistical_alternative_fit["loadings"] if statistical_alternative_fit is not None else pd.DataFrame()),
        "Alternative_Factor_Corr": (statistical_alternative_fit["phi"] if statistical_alternative_fit is not None else pd.DataFrame()),
        "Alternative_Crosswalk": alternative_crosswalk,
        "Alternative_Merge_Candidates": merge_candidates,
        "Alternative_Reliability": alternative_reliability,
        "Alternative_Validity": alternative_validity,
        "Alternative_HTMT": alternative_htmt,
        "Iteration_Log": refinement["iteration_log"],
        "Deleted_Items": refinement["deleted_items"],
        "Alternative_Deletions": refinement["alternative_deletions"],
        "Final_Pattern_Matrix": final_fit["loadings"],
        "Final_Structure_Matrix": final_fit["structure"],
        "Final_Factor_Correlations": final_fit["phi"],
        "Final_Item_Diagnostics": final_item_diagnostics,
        "Final_Residual_Matrix": final_fit["residual"],
        "Final_Method_Sensitivity": sensitivity,
        "Reliability": reliability,
        "Validity": validity,
        "HTMT": htmt,
        "Bootstrap_Stability": stability,
        "CFA_ESEM_Feasibility": pd.DataFrame(flatten_dict(feasibility)),
    }
    notes = [
        "主報告只放必要結論；本檔保留完整中間結果、替代方法與刪題分支。",
        f"主分析相關矩陣：{primary_method}",
        "Pattern matrix為斜交轉軸下的主要解釋表；Structure matrix僅作輔助。",
        "刪題為保守自動程序，仍須由研究者確認題目內容效度。",
    ] + analysis["input_messages"]
    write_side_excel(out / "Questionnaire_quality_details.xlsx", side_sheets, notes)

    (out / "paper_ready_summary.txt").write_text(paper_summary, encoding="utf-8")
    make_docx_report(out / "Questionnaire_quality_report.docx", analysis, figures)
    make_html_report(out / "Questionnaire_quality_report.html", analysis, figures)
    (out / "analysis_manifest.json").write_text(
        json.dumps({
            "input_file": analysis["input_file"],
            "primary_correlation_method": primary_method,
            "final_items": final_items,
            "final_factors": final_fit["metrics"]["k"],
            "final_rotation": final_fit["metrics"]["rotation"],
            "planned_factors": expected_k,
            "statistical_alternative_factors": alternative_k,
            "deleted_items": refinement["deleted_items"].to_dict(orient="records"),
            "outputs": ["Questionnaire_quality_report.docx", "Questionnaire_quality_report.html", "Questionnaire_quality_details.xlsx", "paper_ready_summary.txt", "CFA_lavaan.R", "ESEM_lavaan.R"],
        }, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )
    return analysis


def main() -> int:
    parser = argparse.ArgumentParser(description="Assess Likert questionnaire quality, measurement structure, reliability, and preliminary validity evidence.")
    parser.add_argument("input_excel", help="Path to the Excel questionnaire data file")
    parser.add_argument("--config", default=None, help="Optional YAML configuration file")
    parser.add_argument("--output", default="questionnaire_quality_output", help="Output directory")
    args = parser.parse_args()
    try:
        analysis = run_analysis(args.input_excel, args.output, args.config)
        print("Questionnaire quality assessment completed.")
        print(f"Primary correlation method: {analysis['primary_correlation_method']}")
        print(f"Final items ({len(analysis['final_items'])}): {', '.join(analysis['final_items'])}")
        print(f"Reports saved to: {Path(args.output).resolve()}")
        return 0
    except Exception as exc:
        print(f"ERROR: {type(exc).__name__}: {exc}", file=sys.stderr)
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
